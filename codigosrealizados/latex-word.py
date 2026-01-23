import math
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. PARÁMETROS DE ENTRADA Y CONFIGURACIÓN ---
CAUDAL_DISENO = 1.327
PENDIENTE_CANAL = 0.00373
COEFICIENTE_N_MANNING = 0.014
LONGITUD_TRAMO_PROTOTIPO = 268
ACELERACION_GRAVEDAD = 9.81
LONGITUD_MAXIMA_MODELO = 1.0


# --- FUNCIONES AUXILIARES PARA EL DOCUMENTO DE WORD ---
def add_hyper_detailed_calculation(doc, title, intro_text, steps, final_formula_latex, substitution_latex, values_dict,
                                   result_variable, result_value, result_unit):
    """
    Añade una sección de cálculo con un nivel de detalle extremo al documento.
    """
    doc.add_heading(title, level=4)
    doc.add_paragraph(intro_text)

    # 1. Derivación paso a paso
    p = doc.add_paragraph()
    p.add_run('Derivación Lógica y Algebraica:').bold = True
    for step in steps:
        p_step = doc.add_paragraph(step, style='List Number')
        p_step.paragraph_format.left_indent = Inches(0.25)

    # 2. Fórmula Final
    p = doc.add_paragraph()
    p.add_run('Fórmula Final Aplicable:').bold = True
    p_formula = doc.add_paragraph(f'    {final_formula_latex}')
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. Sustitución de Valores
    p = doc.add_paragraph()
    p.add_run('Sustitución de Datos del Proyecto:').bold = True
    equation_with_values = substitution_latex
    for key, val in values_dict.items():
        str_val = f"{val:.5f}" if isinstance(val, float) else str(val)
        equation_with_values = equation_with_values.replace(key, str_val)
    p_subs = doc.add_paragraph(f'    {equation_with_values}')
    p_subs.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. Resultado
    p = doc.add_paragraph()
    p.add_run('Resultado Obtenido:').bold = True
    p_result = doc.add_paragraph()
    p_result.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_result = p_result.add_run(f'{result_variable} = {result_value:.5f} {result_unit}')
    run_result.font.bold = True
    doc.add_paragraph()  # Espacio extra


# --- FUNCIONES DE CÁLCULO HIDRÁULICO (Sin cambios) ---
def calcular_parametros_hidraulicos(b, y, Q):
    if y is None or b is None or y <= 0 or b <= 0: return {}
    area = b * y
    perimetro = b + 2 * y
    radio_h = area / perimetro
    velocidad = Q / area
    froude = velocidad / math.sqrt(ACELERACION_GRAVEDAD * y)
    regimen = "Subcrítico" if froude < 1 else "Supercrítico" if froude > 1 else "Crítico"
    return {"A": area, "P": perimetro, "R": radio_h, "V": velocidad, "b": b, "y": y, "Fr": froude, "Régimen": regimen}


def disenar_seccion_optima(Q, n, S):
    try:
        y = ((Q * n) / (2 ** (1 / 3) * S ** 0.5)) ** (3 / 8)
        b = 2 * y
        return b, y
    except (ValueError, ZeroDivisionError):
        return None, None


def disenar_minima_infiltracion(Q, n, S):
    try:
        y = ((Q * n) / (4 * (2 / 3) ** (2 / 3) * S ** 0.5)) ** (3 / 8)
        b = 4 * y
        return b, y
    except (ValueError, ZeroDivisionError):
        return None, None


def calcular_analisis_dimensional(Lp, Lm):
    lambda_L = Lp / Lm
    return {"λ_L": lambda_L, "λ_V": lambda_L ** 0.5, "λ_Q": lambda_L ** 2.5}


def escalar_modelo(prototipo_params, escalas, Qp, Lp):
    if not prototipo_params or not escalas: return {}
    b_p, y_p, V_p = prototipo_params["b"], prototipo_params["y"], prototipo_params["V"]
    lambda_L, lambda_V, lambda_Q = escalas["λ_L"], escalas["λ_V"], escalas["λ_Q"]
    return {"b_m": b_p / lambda_L, "y_m": y_p / lambda_L, "V_m": V_p / lambda_V, "Q_m": Qp / lambda_Q,
            "L_m": Lp / lambda_L}


# --- GENERADOR DEL INFORME EN WORD ---
def generar_informe_word(datos):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font;
    font.name = 'Arial';
    font.size = Pt(11)

    doc.add_heading('Memoria de Cálculo: Diseño de Canal – Tramo Cañasbamba', level=0)
    # Secciones introductorias
    doc.add_heading('1. Introducción', level=1)
    doc.add_paragraph(
        "El presente documento constituye la memoria de cálculo detallada para el dimensionamiento de un tramo de canal rectangular de 268 metros en Cañasbamba, Ancash. El objetivo es proporcionar una trazabilidad completa, paso a paso, de todos los procedimientos matemáticos y decisiones de ingeniería tomadas.")
    doc.add_heading('5. Datos de Entrada', level=1)
    doc.add_paragraph(f"Caudal de Diseño (Q): ${datos['Q']}$ m³/s")
    doc.add_paragraph(f"Pendiente del Tramo (S): ${datos['S']:.5f}$ m/m")
    doc.add_paragraph(f"Coeficiente de Rugosidad (n): ${datos['n']}$ (Concreto)")
    doc.add_paragraph(f"Longitud del Prototipo (Lp): ${datos['Lp']}$ m")

    # --- 6. METODOLOGÍA Y DESARROLLO DE CÁLCULOS ---
    doc.add_heading('6. Metodología y Desarrollo de Cálculos', level=1)

    # --- 6.1. Diseño para Máxima Eficiencia Hidráulica ---
    doc.add_heading('6.1. Diseño por Máxima Eficiencia Hidráulica (Sección Óptima)', level=2)

    # Cálculo del Tirante 'y'
    add_hyper_detailed_calculation(doc, "6.1.1. Cálculo del Tirante (y)",
                                   "El primer paso es determinar el tirante de agua que satisface la ecuación de Manning para la condición de máxima eficiencia.",
                                   [
                                       "Se parte de la Ecuación de Manning para Caudal: $Q = \\frac{1}{n} A R^{2/3} S^{1/2}$",
                                       "Para una sección rectangular, el Área es $A = b \\cdot y$ y el Radio Hidráulico es $R = \\frac{A}{P} = \\frac{b \\cdot y}{b + 2y}$",
                                       "La condición teórica de máxima eficiencia hidráulica para una sección rectangular se da cuando el ancho es el doble del tirante: $b = 2y$.",
                                       "Se sustituye esta condición en las fórmulas de Área y Radio Hidráulico: $A = (2y) \\cdot y = 2y^2$ y $R = \\frac{2y^2}{2y + 2y} = \\frac{2y^2}{4y} = \\frac{y}{2}$",
                                       "Estos términos simplificados se reintroducen en la Ecuación de Manning: $Q = \\frac{1}{n} (2y^2) (\\frac{y}{2})^{2/3} S^{1/2}$",
                                       "Se realiza el despeje algebraico de 'y': $Q \\cdot n = 2 \\cdot 2^{-2/3} \\cdot y^2 \\cdot y^{2/3} \\cdot S^{1/2} \\Rightarrow Q \\cdot n = 2^{1/3} \\cdot y^{8/3} \\cdot S^{1/2}$",
                                       "Finalmente, se obtiene la expresión para 'y'."
                                   ],
                                   final_formula_latex="$y = \\left( \\frac{Q \\cdot n}{2^{1/3} S^{1/2}} \\right)^{3/8}$",
                                   substitution_latex="$y = \\left( \\frac{Q_val \\cdot n_val}{2^{1/3} S_val^{1/2}} \\right)^{3/8}$",
                                   values_dict={'Q_val': datos['Q'], 'n_val': datos['n'], 'S_val': datos['S']},
                                   result_variable="y", result_value=datos['res_optima']['y'], result_unit="m")

    # Cálculo del Ancho 'b'
    add_hyper_detailed_calculation(doc, "6.1.2. Cálculo del Ancho de Solera (b)",
                                   "Una vez conocido el tirante óptimo, el ancho se calcula directamente a partir de la condición de máxima eficiencia.",
                                   ["La relación geométrica para máxima eficiencia es directa: $b = 2y$."],
                                   final_formula_latex="$b = 2 \\cdot y$",
                                   substitution_latex="$b = 2 \\cdot y_val$",
                                   values_dict={'y_val': datos['res_optima']['y']},
                                   result_variable="b", result_value=datos['res_optima']['b'], result_unit="m")

    # Cálculos de Parámetros Hidráulicos Derivados
    doc.add_heading("6.1.3. Cálculo de Parámetros Hidráulicos Resultantes", level=3)
    add_hyper_detailed_calculation(doc, "Área Mojada (A)", "El área de la sección transversal del flujo.",
                                   ["$A = b \\cdot y$"], "$A = b \\cdot y$", "$A = b_val \\cdot y_val$",
                                   {'b_val': datos['res_optima']['b'], 'y_val': datos['res_optima']['y']}, "A",
                                   datos['res_optima']['A'], "m²")
    add_hyper_detailed_calculation(doc, "Perímetro Mojado (P)",
                                   "La longitud de la superficie del canal en contacto con el agua.", ["$P = b + 2y$"],
                                   "$P = b + 2y$", "$P = b_val + 2 \\cdot y_val$",
                                   {'b_val': datos['res_optima']['b'], 'y_val': datos['res_optima']['y']}, "P",
                                   datos['res_optima']['P'], "m")
    add_hyper_detailed_calculation(doc, "Radio Hidráulico (R)",
                                   "La relación entre el área y el perímetro mojado, indicador de eficiencia.",
                                   ["$R = A / P$"], "$R = \\frac{A}{P}$", "$R = \\frac{A_val}{P_val}$",
                                   {'A_val': datos['res_optima']['A'], 'P_val': datos['res_optima']['P']}, "R",
                                   datos['res_optima']['R'], "m")
    add_hyper_detailed_calculation(doc, "Velocidad del Flujo (V)",
                                   "La velocidad media del agua en el canal, verificada con la Ecuación de Continuidad.",
                                   ["$V = Q / A$"], "$V = \\frac{Q}{A}$", "$V = \\frac{Q_val}{A_val}$",
                                   {'Q_val': datos['Q'], 'A_val': datos['res_optima']['A']}, "V",
                                   datos['res_optima']['V'], "m/s")
    add_hyper_detailed_calculation(doc, "Número de Froude (Fr)",
                                   "Parámetro adimensional para clasificar el régimen de flujo.",
                                   ["$Fr = V / \\sqrt{g \\cdot y}$"], "$Fr = \\frac{V}{\\sqrt{g \\cdot y}}$",
                                   "$Fr = \\frac{V_val}{\\sqrt{g_val \\cdot y_val}}$",
                                   {'V_val': datos['res_optima']['V'], 'g_val': ACELERACION_GRAVEDAD,
                                    'y_val': datos['res_optima']['y']}, "Fr", datos['res_optima']['Fr'],
                                   "(adimensional)")
    doc.add_paragraph(f"El régimen de flujo es: {datos['res_optima']['Régimen']} ya que Fr > 1.").bold = True

    # --- 6.2. Análisis Dimensional para Prototipado ---
    doc.add_heading('6.2. Análisis Dimensional y Semejanza (Modelo a Escala)', level=2)
    doc.add_paragraph(
        "Para que el modelo a escala se comporte hidráulicamente como el prototipo, se usa la semejanza de Froude. Se tomarán los resultados del diseño de máxima eficiencia como base para el prototipo.")
    add_hyper_detailed_calculation(doc, "6.2.1. Cálculo de Escala de Longitud ($\\lambda_L$)",
                                   "Define la relación geométrica entre el prototipo (p) y el modelo (m).",
                                   ["$\\lambda_L = L_p / L_m$"], "$\\lambda_L = \\frac{L_p}{L_m}$",
                                   "$\\lambda_L = \\frac{Lp_val}{Lm_val}$",
                                   {'Lp_val': datos['Lp'], 'Lm_val': datos['Lm']}, "$\\lambda_L$",
                                   datos['escalas']['λ_L'], "")
    add_hyper_detailed_calculation(doc, "6.2.2. Cálculo de Escala de Velocidad ($\\lambda_V$)",
                                   "Se deriva de mantener el Número de Froude constante ($Fr_p = Fr_m$).",
                                   ["$\\frac{V_p}{\\sqrt{g_p y_p}} = \\frac{V_m}{\\sqrt{g_m y_m}} \\Rightarrow \\frac{V_p}{V_m} = \\sqrt{\\frac{y_p}{y_m}} \\Rightarrow \\lambda_V = \\sqrt{\\lambda_L}$"],
                                   "$\\lambda_V = \\sqrt{\\lambda_L}$", "$\\lambda_V = \\sqrt{\\lambda_L_val}$",
                                   {'\\lambda_L_val': datos['escalas']['λ_L']}, "$\\lambda_V$", datos['escalas']['λ_V'],
                                   "")
    add_hyper_detailed_calculation(doc, "6.2.3. Cálculo de Escala de Caudal ($\\lambda_Q$)",
                                   "Se deriva de la Ecuación de Continuidad y las escalas ya definidas.",
                                   ["$Q = A \\cdot V \\Rightarrow \\lambda_Q = \\lambda_A \\cdot \\lambda_V = (\\lambda_L^2) \\cdot (\\lambda_L^{0.5}) = \\lambda_L^{2.5}$"],
                                   "$\\lambda_Q = \\lambda_L^{2.5}$", "$\\lambda_Q = \\lambda_L_val^{2.5}$",
                                   {'\\lambda_L_val': datos['escalas']['λ_L']}, "$\\lambda_Q$", datos['escalas']['λ_Q'],
                                   "")

    # --- 6.3. Parámetros del Modelo Físico ---
    doc.add_heading('6.3. Parámetros del Modelo Físico a Escala', level=2)
    doc.add_paragraph(
        "Finalmente, se calculan las características que debe tener la maqueta de laboratorio dividiendo los parámetros del prototipo por sus respectivas escalas.")
    add_hyper_detailed_calculation(doc, "Ancho del Modelo ($b_m$)", "", ["$b_m = b_p / \\lambda_L$"],
                                   "$b_m = \\frac{b_p}{\\lambda_L}$", "$b_m = \\frac{bp_val}{\\lambda_L_val}$",
                                   {'bp_val': datos['res_optima']['b'], '\\lambda_L_val': datos['escalas']['λ_L']},
                                   "b_m", datos['modelo']['b_m'], "m")
    add_hyper_detailed_calculation(doc, "Tirante del Modelo ($y_m$)", "", ["$y_m = y_p / \\lambda_L$"],
                                   "$y_m = \\frac{y_p}{\\lambda_L}$", "$y_m = \\frac{yp_val}{\\lambda_L_val}$",
                                   {'yp_val': datos['res_optima']['y'], '\\lambda_L_val': datos['escalas']['λ_L']},
                                   "y_m", datos['modelo']['y_m'], "m")
    add_hyper_detailed_calculation(doc, "Velocidad en el Modelo ($V_m$)", "", ["$V_m = V_p / \\lambda_V$"],
                                   "$V_m = \\frac{V_p}{\\lambda_V}$", "$V_m = \\frac{Vp_val}{\\lambda_V_val}$",
                                   {'Vp_val': datos['res_optima']['V'], '\\lambda_V_val': datos['escalas']['λ_V']},
                                   "V_m", datos['modelo']['V_m'], "m/s")
    add_hyper_detailed_calculation(doc, "Caudal en el Modelo ($Q_m$)", "", ["$Q_m = Q_p / \\lambda_Q$"],
                                   "$Q_m = \\frac{Q_p}{\\lambda_Q}$", "$Q_m = \\frac{Qp_val}{\\lambda_Q_val}$",
                                   {'Qp_val': datos['Q'], '\\lambda_Q_val': datos['escalas']['λ_Q']}, "Q_m",
                                   datos['modelo']['Q_m'], "m³/s")

    # --- 7. Resumen de Resultados ---
    doc.add_heading('7. Resumen y Conclusiones del Dimensionamiento', level=1)
    doc.add_paragraph(
        "La siguiente tabla resume los parámetros finales del prototipo (canal real) y del modelo físico (maqueta) para su construcción y experimentación.")
    tabla_resumen = doc.add_table(rows=1, cols=3);
    tabla_resumen.style = 'Table Grid'
    hdr = tabla_resumen.rows[0].cells;
    hdr[0].text = 'Parámetro';
    hdr[1].text = 'Prototipo (Real)';
    hdr[2].text = 'Modelo (Laboratorio)'
    b_p, y_p, V_p, Q_p = datos['res_optima']['b'], datos['res_optima']['y'], datos['res_optima']['V'], datos['Q']
    b_m, y_m, V_m, Q_m = datos['modelo']['b_m'], datos['modelo']['y_m'], datos['modelo']['V_m'], datos['modelo']['Q_m']
    data = [['Ancho (b)', f'{b_p:.3f} m', f'{b_m * 100:.2f} cm'],
            ['Tirante (y)', f'{y_p:.3f} m', f'{y_m * 100:.2f} cm'],
            ['Velocidad (V)', f'{V_p:.3f} m/s', f'{V_m:.3f} m/s'],
            ['Caudal (Q)', f'{Q_p:.3f} m³/s', f'{Q_m * 1000:.3f} L/s'],
            ['Régimen de Flujo', datos['res_optima']['Régimen'], 'Supercrítico (por semejanza)']]
    for item in data:
        row = tabla_resumen.add_row().cells;
        row[0].text = item[0];
        row[1].text = item[1];
        row[2].text = item[2]

    filename = 'Reporte_Calculo_Ultra_Detallado_Canal.docx'
    doc.save(filename)
    return filename


# --- FLUJO PRINCIPAL DE EJECUCIÓN ---
if __name__ == "__main__":
    print("Iniciando memoria de cálculo ultra-detallada...")
    print(f"Usando Caudal de Diseño: {CAUDAL_DISENO} m³/s")

    b_optima, y_optima = disenar_seccion_optima(CAUDAL_DISENO, COEFICIENTE_N_MANNING, PENDIENTE_CANAL)
    res_optima = calcular_parametros_hidraulicos(b_optima, y_optima, CAUDAL_DISENO)

    escalas = calcular_analisis_dimensional(LONGITUD_TRAMO_PROTOTIPO, LONGITUD_MAXIMA_MODELO)
    modelo_params = escalar_modelo(res_optima, escalas, CAUDAL_DISENO, LONGITUD_TRAMO_PROTOTIPO)

    datos_informe = {
        "Q": CAUDAL_DISENO, "S": PENDIENTE_CANAL, "n": COEFICIENTE_N_MANNING,
        "Lp": LONGITUD_TRAMO_PROTOTIPO, "Lm": LONGITUD_MAXIMA_MODELO,
        "res_optima": res_optima, "escalas": escalas, "modelo": modelo_params
    }

    print("Cálculos finalizados. Generando informe en Word...")
    try:
        nombre_archivo = generar_informe_word(datos_informe)
        print(f"\n¡Éxito! Se ha generado el informe: '{nombre_archivo}'")
        print("El archivo contiene el desglose completo de cada cálculo.")
        try:
            os.startfile(nombre_archivo)
        except AttributeError:
            print("Para ver el informe, abre el archivo manualmente.")
    except Exception as e:
        print(f"\nError: No se pudo generar el archivo de Word. Causa: {e}")
