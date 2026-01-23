import tkinter as tk
from tkinter import ttk, messagebox
import numpy as np
from scipy.optimize import brentq
import logging
import sys
import math
import datetime
import os
import io

# --- Importaciones de Ingeniería ---
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
except ImportError:
    print("CRÍTICO: Falta 'python-docx'. Ejecuta: pip install python-docx")
    sys.exit(1)

try:
    import matplotlib

    matplotlib.use('TkAgg')
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    from matplotlib.figure import Figure
    import matplotlib.pyplot as plt

    plt.rcParams["font.family"] = "serif"
    plt.rcParams["mathtext.fontset"] = "cm"
except ImportError:
    print("CRÍTICO: Falta 'matplotlib'. Ejecuta: pip install matplotlib")
    sys.exit(1)

logging.basicConfig(level=logging.INFO, format='%(levelname)s - %(message)s')


class HydraulicMath:
    @staticmethod
    def manning_solver(Q, b, n, S):
        """Resuelve el tirante normal (y_n)"""

        def func(y):
            if y <= 0: return 1e9
            A = b * y
            P = b + 2 * y
            return ((1 / n) * A * ((A / P) ** (2 / 3)) * (S ** 0.5)) - Q

        try:
            return brentq(func, 1e-4, 10.0)
        except:
            return 0.0

    @staticmethod
    def conjugate_depth(y1, Fr1):
        """Calcula y2 para resalto hidráulico"""
        return (y1 / 2.0) * (math.sqrt(1 + 8 * (Fr1 ** 2)) - 1)


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("EMAIRC Suite - Proyecto Río Santa (DD1)")
        self.geometry("1300x900")
        self.state('zoomed')

        style = ttk.Style()
        style.theme_use('clam')
        style.configure("TLabel", font=("Segoe UI", 10))
        style.configure("Title.TLabel", font=("Segoe UI", 14, "bold"),
                        foreground="#8B0000")  # Rojo oscuro estilo académico
        style.configure("Group.TLabelframe.Label", font=("Segoe UI", 10, "bold"))

        # --- VARIABLES PRECARGADAS (DATOS RÍO SANTA) ---
        self.v = {
            # 1. Datos Lab (Para llegar a Qm ≈ 0.177 L/s)
            "vol": tk.DoubleVar(value=1.77),  # Litros
            "t": tk.DoubleVar(value=10.0),  # Segundos
            "bm": tk.DoubleVar(value=0.10),  # 3m / 30 = 0.10m
            "nm": tk.DoubleVar(value=0.010),  # Vidrio/Acrílico
            "Sm": tk.DoubleVar(value=0.015),  # 1.5% (Dato Real)

            # 2. Escala
            "Lr": tk.DoubleVar(value=30.0),  # Escala 1:30

            # 3. Prototipo (Validación)
            "np_real": tk.DoubleVar(value=0.035),  # Manning Río Santa

            # 4. Resultados Validación Numérica (Datos de tu Excel)
            "y_hec": tk.DoubleVar(value=0.69),
            "v_hec": tk.DoubleVar(value=4.21),
            "y_iber": tk.DoubleVar(value=0.72),
            "v_iber": tk.DoubleVar(value=4.04),
            "y_foam": tk.DoubleVar(value=0.74),
            "v_foam": tk.DoubleVar(value=3.92),

            # 5. Compuerta y Resalto
            "Cd": tk.DoubleVar(value=0.61),
            "ap": tk.DoubleVar(value=0.10),  # Abertura tentativa en prototipo

            # Resultados Calculados
            "Qm": tk.DoubleVar(), "Qp": tk.DoubleVar(),
            "yn_p": tk.DoubleVar(), "Fr_p": tk.DoubleVar(),
            "y1_jump": tk.DoubleVar(), "y2_jump": tk.DoubleVar()
        }

        self.setup_ui()

    def setup_ui(self):
        main_pane = ttk.PanedWindow(self, orient=tk.HORIZONTAL)
        main_pane.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        scroll_frame = ttk.Frame(main_pane)
        # (Para simplicidad usamos frame directo, si es muy grande necesitaría canvas)

        left_frame = ttk.Frame(main_pane)
        right_frame = ttk.Frame(main_pane)
        main_pane.add(left_frame, weight=1)
        main_pane.add(right_frame, weight=2)

        # --- PANEL IZQUIERDO ---
        ttk.Label(left_frame, text="Panel de Control: Río Santa", style="Title.TLabel").pack(pady=10)

        # G1: Datos Experimentales y Escala
        lf1 = ttk.LabelFrame(left_frame, text="1. Modelo Físico (Escala 1:30)", style="Group.TLabelframe")
        lf1.pack(fill=tk.X, padx=5, pady=5)
        self.grid_entry(lf1, "Volumen (L):", self.v["vol"], 0)
        self.grid_entry(lf1, "Tiempo (s):", self.v["t"], 1)
        self.grid_entry(lf1, "Ancho Modelo (m):", self.v["bm"], 2)
        self.grid_entry(lf1, "Escala (Lr):", self.v["Lr"], 3)
        self.grid_entry(lf1, "Pendiente (m/m):", self.v["Sm"], 4)

        # G2: Validación Numérica
        lf2 = ttk.LabelFrame(left_frame, text="2. Validación Numérica (Comparativa)", style="Group.TLabelframe")
        lf2.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(lf2, text="HEC-RAS (y, v):").grid(row=0, column=0)
        ttk.Entry(lf2, textvariable=self.v["y_hec"], width=6).grid(row=0, column=1)
        ttk.Entry(lf2, textvariable=self.v["v_hec"], width=6).grid(row=0, column=2)

        ttk.Label(lf2, text="IBER (y, v):").grid(row=1, column=0)
        ttk.Entry(lf2, textvariable=self.v["y_iber"], width=6).grid(row=1, column=1)
        ttk.Entry(lf2, textvariable=self.v["v_iber"], width=6).grid(row=1, column=2)

        ttk.Label(lf2, text="OpenFOAM (y, v):").grid(row=2, column=0)
        ttk.Entry(lf2, textvariable=self.v["y_foam"], width=6).grid(row=2, column=1)
        ttk.Entry(lf2, textvariable=self.v["v_foam"], width=6).grid(row=2, column=2)

        # G3: Prototipo y Estructura
        lf3 = ttk.LabelFrame(left_frame, text="3. Prototipo y Compuerta", style="Group.TLabelframe")
        lf3.pack(fill=tk.X, padx=5, pady=5)
        self.grid_entry(lf3, "Manning Real (np):", self.v["np_real"], 0)
        self.grid_entry(lf3, "Abertura Comp. (m):", self.v["ap"], 1)
        self.grid_entry(lf3, "Coef. Descarga (Cd):", self.v["Cd"], 2)

        # Botones
        ttk.Button(left_frame, text="CALCULAR TODO", command=self.calculate).pack(fill=tk.X, pady=10)
        ttk.Button(left_frame, text="GENERAR INFORME DD1 (WORD)", command=self.generate_report).pack(fill=tk.X, pady=5)

        self.lbl_res = ttk.Label(left_frame, text="...", font=("Consolas", 9), foreground="#333")
        self.lbl_res.pack(pady=10)

        # --- PANEL DERECHO (GRÁFICO) ---
        self.fig = Figure(figsize=(5, 4), dpi=100)
        self.ax = self.fig.add_subplot(111)
        self.canvas = FigureCanvasTkAgg(self.fig, master=right_frame)
        self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    def grid_entry(self, parent, text, var, row):
        ttk.Label(parent, text=text).grid(row=row, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(parent, textvariable=var, width=10).grid(row=row, column=1, sticky="e", padx=5, pady=2)

    def calculate(self):
        try:
            # 1. Experimental
            Qm = (self.v["vol"].get() / 1000) / self.v["t"].get()
            self.v["Qm"].set(Qm)

            # 2. Escalamiento
            Lr = self.v["Lr"].get()
            Qp = Qm * (Lr ** 2.5)
            self.v["Qp"].set(Qp)

            bp = self.v["bm"].get() * Lr
            np_real = self.v["np_real"].get()
            Sm = self.v["Sm"].get()

            # 3. Tirante Normal (Prototipo)
            yn = HydraulicMath.manning_solver(Qp, bp, np_real, Sm)
            self.v["yn_p"].set(yn)
            Vn = Qp / (bp * yn)
            Fr = Vn / math.sqrt(9.81 * yn)
            self.v["Fr_p"].set(Fr)

            # 4. Compuerta y Resalto
            a = self.v["ap"].get()
            Cc = 0.61
            y1 = Cc * a
            V1 = Qp / (bp * y1)
            Fr1 = V1 / math.sqrt(9.81 * y1)
            y2 = HydraulicMath.conjugate_depth(y1, Fr1)

            self.v["y1_jump"].set(y1)
            self.v["y2_jump"].set(y2)

            self.lbl_res.config(text=f"Qp: {Qp:.3f} m³/s | yn: {yn:.3f} m | Fr: {Fr:.3f}\n"
                                     f"Resalto: y1={y1:.2f}m -> y2={y2:.2f}m")

            self.plot(bp, yn, y1, y2)

        except Exception as e:
            messagebox.showerror("Error", str(e))

    def plot(self, b, yn, y1, y2):
        self.ax.clear()
        self.ax.set_title("Perfil Hidráulico: Río Santa")
        self.ax.set_ylabel("Tirante (m)")

        # Puntos esquemáticos
        x = [0, 2, 8, 15]
        y_vals = [y1, y1, y2, yn]  # Supercrítico -> Resalto -> Normal

        self.ax.plot(x, y_vals, 'b-o', label="Superficie Libre")
        self.ax.axhline(y=yn, color='g', linestyle='--', label=f"Tirante Normal ({yn:.2f}m)")
        self.ax.axhline(y=y2, color='r', linestyle=':', label=f"Conjugado ({y2:.2f}m)")

        self.ax.legend()
        self.ax.grid(True)
        self.canvas.draw()

    def generate_report(self):
        # Renderizado de ecuaciones LaTeX (Backend Agg)
        import matplotlib.pyplot as plt_render
        plt_render.switch_backend('Agg')

        def latex_to_img(latex_str):
            fig = plt_render.figure(figsize=(4, 0.6))
            fig.text(0.5, 0.5, f"${latex_str}$", fontsize=12, ha='center', va='center')
            buf = io.BytesIO()
            plt_render.savefig(buf, format='png', dpi=150, transparent=True, bbox_inches='tight')
            plt_render.close(fig)
            buf.seek(0)
            return buf

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5

        # --- CONTENIDO ESPECÍFICO RÍO SANTA ---
        doc.add_heading('INFORME DD1: RÍO SANTA (CUENCA ALTA)', 0)
        doc.add_paragraph(f"Fecha: {datetime.date.today()}\n")

        # Tabla Comparativa (Lo más importante de tu Excel)
        doc.add_heading('6.5 Resultados y Validación Numérica', level=2)
        doc.add_paragraph(
            "A continuación se presenta la comparación entre los datos experimentales escalados al prototipo y las simulaciones numéricas.")

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Método'
        hdr[1].text = 'Tirante (m)'
        hdr[2].text = 'Velocidad (m/s)'
        hdr[3].text = 'Error Relativo (%)'

        # Datos Experimentales (Prototipo)
        yn_exp = self.v["yn_p"].get()
        vn_exp = self.v["Qp"].get() / ((self.v["bm"].get() * self.v["Lr"].get()) * yn_exp)

        data_rows = [
            ("Experimental (Físico)", f"{yn_exp:.3f}", f"{vn_exp:.3f}", "-"),
            ("HEC-RAS", f"{self.v['y_hec'].get()}", f"{self.v['v_hec'].get()}",
             f"{abs(yn_exp - self.v['y_hec'].get()) / yn_exp * 100:.1f}%"),
            ("IBER", f"{self.v['y_iber'].get()}", f"{self.v['v_iber'].get()}",
             f"{abs(yn_exp - self.v['y_iber'].get()) / yn_exp * 100:.1f}%"),
            ("OpenFOAM", f"{self.v['y_foam'].get()}", f"{self.v['v_foam'].get()}",
             f"{abs(yn_exp - self.v['y_foam'].get()) / yn_exp * 100:.1f}%")
        ]

        for m, y, v, e in data_rows:
            r = table.add_row().cells
            r[0].text = m;
            r[1].text = y;
            r[2].text = v;
            r[3].text = e

        # Sección de Resalto
        doc.add_heading('7.2 Análisis del Resalto Hidráulico', level=2)
        doc.add_paragraph(
            "Se evaluó la formación del resalto hidráulico posterior a la compuerta. Los tirantes conjugados teóricos son:")

        img = latex_to_img(r"\frac{y_2}{y_1} = \frac{1}{2} (\sqrt{1 + 8 Fr_1^2} - 1)")
        doc.add_paragraph().add_run().add_picture(img)

        p = doc.add_paragraph()
        p.add_run(f"Tirante Supercrítico (y1): {self.v['y1_jump'].get():.3f} m\n")
        p.add_run(f"Tirante Subcrítico Conjugado (y2): {self.v['y2_jump'].get():.3f} m\n")

        fname = f"Informe_RioSanta_DD1_{datetime.datetime.now().strftime('%H%M%S')}.docx"
        doc.save(fname)
        plt_render.switch_backend('TkAgg')
        messagebox.showinfo("Éxito", f"Informe Río Santa generado:\n{fname}")


if __name__ == "__main__":
    app = App()
    app.mainloop()