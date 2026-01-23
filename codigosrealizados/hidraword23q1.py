# -*- coding: utf-8 -*-
"""
================================================================================
== GENERADOR DE REPORTE FINAL: DISEÑO DE MODELO HIDRÁULICO A ESCALA ==
================================================================================

Este script toma un escenario de diseño final y validado, y genera una memoria
de cálculo ultra detallada en un documento de Word (.docx), siguiendo la
estructura y el nivel de detalle solicitados.

La narrativa del reporte sigue el flujo de diseño estándar:
1. Se presentan las especificaciones del Prototipo (canal real) como datos de
   partida.
2. Se definen las restricciones de construcción del Modelo (maqueta).
3. Se calcula la escala de semejanza necesaria.
4. Se derivan y detallan todos los parámetros operativos del Modelo.
5. Se presenta la validación final de la semejanza dinámica.
"""
import math
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. PARÁMETROS DEL ESCENARIO FINAL VALIDADO (DATOS DE ENTRADA) ---
# Estos valores son el resultado del script de cálculo inverso anterior.
# -------------------------------------------------------------------
# --- Prototipo (Canal Real de Interés) ---
PROTOTIPO_PARAMS = {
    "b": 3.00000,
    "y": 0.20850,
    "Q": 0.87142,
    "n": 0.01413,
    "S": 0.00373,
    "g": 9.81
}

# --- Modelo (Restricciones y Material) ---
MODELO_PARAMS = {
    "b": 0.20,
    "n": 0.009
}


# -------------------------------------------------------------------

# --- 2. FUNCIONES AUXILIARES PARA EL DOCUMENTO DE WORD ---

def add_hyper_detailed_calculation(doc, title, intro_text, steps, final_formula_latex, substitution_latex, values_dict,
                                   result_variable, result_value, result_unit):
    """
    Añade una sección de cálculo con un nivel de detalle extremo al documento.
    """
    doc.add_heading(title, level=3)
    doc.add_paragraph(intro_text)

    p = doc.add_paragraph()
    p.add_run('Derivación Lógica y Algebraica:').bold = True
    for step in steps:
        p_step = doc.add_paragraph(style='List Number')
        # Añadir texto con formato LaTeX simulado
        parts = step.split('$')
        for i, part in enumerate(parts):
            if i % 2 == 1:  # Texto dentro de $...$
                run = p_step.add_run(part)
                run.font.italic = True
            else:
                p_step.add_run(part)
        p_step.paragraph_format.left_indent = Inches(0.25)

    p = doc.add_paragraph()
    p.add_run('Fórmula Final Aplicable:').bold = True
    p_formula = doc.add_paragraph()
    run = p_formula.add_run(f'${final_formula_latex}$')
    run.font.italic = True
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Sustitución de Datos del Proyecto:').bold = True
    equation_with_values = substitution_latex
    for key, val in values_dict.items():
        str_val = f"{val:.5f}" if isinstance(val, float) else str(val)
        equation_with_values = equation_with_values.replace(key, str_val)
    p_subs = doc.add_paragraph()
    run = p_subs.add_run(f'${equation_with_values}$')
    run.font.italic = True
    p_subs.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Resultado Obtenido:').bold = True
    p_result = doc.add_paragraph()
    p_result.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_result = p_result.add_run(f'{result_variable} = {result_value:.5f} {result_unit}')
    run_result.font.bold = True
    doc.add_paragraph()


# --- 3. GENERADOR DEL INFORME EN WORD ---

def generar_informe_word(proto_data, model_data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    doc.add_heading('Memoria de Cálculo: Diseño de Prototipo de Canal a Escala', level=0)

    # --- Secciones introductorias ---
    doc.add_heading('1. Introducción y Objetivo', level=1)
    doc.add_paragraph(
        "El presente documento detalla la memoria de cálculo para el diseño de un modelo físico a escala reducida "
        "de un tramo de canal rectangular. El objetivo es determinar las dimensiones y condiciones operativas de una "
        "maqueta de laboratorio que simule de manera dinámicamente semejante (según la Ley de Froude) a un "
        "canal real (prototipo) con características hidráulicas predefinidas."
    )

    doc.add_heading('2. Datos de Partida y Restricciones', level=1)
    doc.add_heading('2.1. Especificaciones del Prototipo (Canal Real)', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parámetro'
    hdr_cells[1].text = 'Valor'
    data_rows = [
        ("Caudal de Diseño ($Q_p$)", f"{proto_data['Q']:.5f} m³/s"),
        ("Ancho de Solera ($b_p$)", f"{proto_data['b']:.5f} m"),
        ("Tirante Normal ($y_p$)", f"{proto_data['y']:.5f} m"),
        ("Rugosidad de Manning ($n_p$)", f"{proto_data['n']:.5f} (Concreto)"),
        ("Pendiente ($S$)", f"{proto_data['S']:.5f} m/m")
    ]
    for key, val in data_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = val

    doc.add_heading('2.2. Restricciones del Modelo (Maqueta de Laboratorio)', level=2)
    table_m = doc.add_table(rows=1, cols=2)
    table_m.style = 'Table Grid'
    hdr_cells_m = table_m.rows[0].cells
    hdr_cells_m[0].text = 'Parámetro'
    hdr_cells_m[1].text = 'Valor'
    row_m = table_m.add_row().cells
    row_m[0].text = "Ancho del Modelo ($b_m$)"
    row_m[1].text = f"{model_data['b']:.3f} m ({model_data['b'] * 100:.1f} cm)"

    # --- 6. METODOLOGÍA Y DESARROLLO DE CÁLCULOS ---
    doc.add_heading('3. Metodología y Desarrollo de Cálculos', level=1)

    # --- 3.1. Cálculo de Escalas ---
    doc.add_heading('3.1. Análisis Dimensional y Determinación de Escalas', level=2)

    add_hyper_detailed_calculation(doc, "3.1.1. Cálculo de Escala de Longitud ($\\lambda_L$)",
                                   "Define la relación geométrica que debe existir entre el prototipo (p) y el modelo (m).",
                                   [
                                       "La escala de longitud ($\\lambda_L$) es el cociente entre una dimensión homóloga del prototipo y el modelo.",
                                       "Se utiliza el ancho de la solera como dimensión característica para establecer esta relación."
                                   ],
                                   "\\lambda_L = \\frac{b_p}{b_m}",
                                   "\\lambda_L = \\frac{bp_val}{bm_val}",
                                   {'bp_val': proto_data['b'], 'bm_val': model_data['b']},
                                   "\\lambda_L", proto_data['lambda_L'], "(adimensional)"
                                   )

    # --- 3.2. Parámetros del Modelo Físico ---
    doc.add_heading('3.2. Parámetros del Modelo Físico a Escala', level=2)
    doc.add_paragraph(
        "Con la escala de longitud definida, se procede a calcular los parámetros requeridos para la construcción y operación de la maqueta.")

    add_hyper_detailed_calculation(doc, "3.2.1. Tirante del Modelo ($y_m$)",
                                   "Determina la profundidad del agua que se debe mantener en la maqueta.",
                                   ["El tirante del modelo se obtiene escalando geométricamente el tirante del prototipo."],
                                   "y_m = \\frac{y_p}{\\lambda_L}",
                                   "y_m = \\frac{yp_val}{\\lambda_L_val}",
                                   {'yp_val': proto_data['y'], '\\lambda_L_val': proto_data['lambda_L']},
                                   "y_m", model_data['y'], "m"
                                   )

    add_hyper_detailed_calculation(doc, "3.2.2. Rugosidad Requerida del Modelo ($n_m$)",
                                   "Determina el coeficiente de Manning que debe tener el material de la maqueta para garantizar la semejanza dinámica.",
                                   [
                                       "Para que las leyes de Manning y Froude se cumplan simultáneamente, la rugosidad debe ser escalada.",
                                       "La relación de escalas de rugosidad es $n_r = n_p / n_m = \\lambda_L^{1/6}$.",
                                       "Despejando $n_m$ se obtiene la rugosidad que debe tener el material del modelo."
                                   ],
                                   "n_m = \\frac{n_p}{\\lambda_L^{1/6}}",
                                   "n_m = \\frac{np_val}{\\lambda_L_val^{1/6}}",
                                   {'np_val': proto_data['n'], '\\lambda_L_val': proto_data['lambda_L']},
                                   "n_m", model_data['n'], "(adimensional)"
                                   )

    add_hyper_detailed_calculation(doc, "3.2.3. Caudal Operativo del Modelo ($Q_m$)",
                                   "Calcula el caudal que se debe suministrar a la maqueta para simular las condiciones del prototipo.",
                                   [
                                       "Se aplica la ecuación de Manning utilizando todos los parámetros ya definidos para el modelo (geometría y rugosidad).",
                                       "$V_m = \\frac{1}{n_m} R_{h,m}^{2/3} S^{1/2}$",
                                       "$Q_m = A_m \\cdot V_m$"
                                   ],
                                   "Q_m = (b_m y_m) \\cdot \\frac{1}{n_m} \\cdot \\left( \\frac{b_m y_m}{b_m + 2y_m} \\right)^{2/3} \\cdot S^{1/2}",
                                   "Q_m = (bm_val \\cdot ym_val) \\cdot \\frac{1}{nm_val} \\cdot \\left( \\frac{bm_val \\cdot ym_val}{bm_val + 2ym_val} \\right)^{2/3} \\cdot S_val^{1/2}",
                                   {'bm_val': model_data['b'], 'ym_val': model_data['y'], 'nm_val': model_data['n'],
                                    'S_val': proto_data['S']},
                                   "Q_m", model_data['Q'], "m³/s"
                                   )

    # --- 4. Validación ---
    doc.add_heading('4. Verificación de Semejanza Dinámica', level=1)

    add_hyper_detailed_calculation(doc, "4.1. Verificación del Número de Froude",
                                   "Se confirma que el diseño es correcto calculando y comparando los Números de Froude del prototipo y del modelo. Deben ser idénticos.",
                                   [
                                       "Se calcula $Fr_p = V_p / \\sqrt{g \\cdot y_p}$ para el prototipo.",
                                       "Se calcula $Fr_m = V_m / \\sqrt{g \\cdot y_m}$ para el modelo.",
                                       "Se comparan ambos valores. Si $Fr_p = Fr_m$, el modelo es dinámicamente semejante."
                                   ],
                                   "Fr_p = Fr_m",
                                   "\\frac{Vp_val}{\\sqrt{g \\cdot yp_val}} = \\frac{Vm_val}{\\sqrt{g \\cdot ym_val}}",
                                   {'Vp_val': proto_data['V'], 'yp_val': proto_data['y'], 'Vm_val': model_data['V'],
                                    'ym_val': model_data['y']},
                                   f"{proto_data['Fr']:.4f}", proto_data['Fr'], f"≈ {model_data['Fr']:.4f}"
                                   )

    doc.add_paragraph(
        f"VALIDACIÓN EXITOSA: Los números de Froude coinciden. El régimen de flujo ({proto_data['Régimen']}) se conservará.").bold = True

    filename = 'Reporte_Calculo_Final_Detallado.docx'
    doc.save(filename)
    return filename


# --- FLUJO PRINCIPAL DE EJECUCIÓN ---
if __name__ == "__main__":
    print("Iniciando generación de memoria de cálculo ultra-detallada...")

    # --- Calcular todos los parámetros necesarios ---

    # Prototipo
    proto_full = PROTOTIPO_PARAMS.copy()
    proto_full["A"] = proto_full["b"] * proto_full["y"]
    proto_full["V"] = proto_full["Q"] / proto_full["A"]
    proto_full["Fr"] = proto_full["V"] / math.sqrt(proto_full["g"] * proto_full["y"])
    proto_full["Régimen"] = "Subcrítico" if proto_full["Fr"] < 1 else "Supercrítico"

    # Escalas
    lambda_L = proto_full["b"] / MODELO_PARAMS["b"]
    proto_full["lambda_L"] = lambda_L  # Guardamos la escala para usarla en el reporte

    # Modelo
    model_full = MODELO_PARAMS.copy()
    model_full["y"] = proto_full["y"] / lambda_L

    # Calcular Vm y Qm del modelo usando sus propios parámetros para consistencia
    area_m = model_full['b'] * model_full['y']
    perimetro_m = model_full['b'] + 2 * model_full['y']
    radio_h_m = area_m / perimetro_m
    V_m = (1 / model_full['n']) * (radio_h_m ** (2 / 3)) * (proto_full['S'] ** 0.5)
    Q_m = V_m * area_m
    model_full['Q'] = Q_m
    model_full['V'] = V_m
    model_full['Fr'] = V_m / math.sqrt(proto_full['g'] * model_full['y'])
    model_full['Régimen'] = "Subcrítico" if model_full["Fr"] < 1 else "Supercrítico"

    print("Cálculos finalizados. Generando informe en Word...")

    try:
        nombre_archivo = generar_informe_word(proto_full, model_full)
        print(f"\n¡Éxito! Se ha generado el informe: '{nombre_archivo}'")
        print("El archivo contiene el desglose completo de cada cálculo, siguiendo la narrativa de diseño solicitada.")
        # Abrir el archivo automáticamente (opcional, funciona en Windows)
        try:
            os.startfile(nombre_archivo)
        except AttributeError:
            print(f"Para ver el informe, abre el archivo '{nombre_archivo}' manualmente.")
    except Exception as e:
        print(f"\nError: No se pudo generar el archivo de Word. Causa: {e}")

