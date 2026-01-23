import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# --- DATOS CENTRALIZADOS ---
# Esta lista de diccionarios contiene todos los cálculos pre-procesados
# para asegurar consistencia total entre todos los archivos.

datos = []

# Constantes
Q = 0.01111
So = 0.01200
b = 0.31
n_base = 0.01
n_pared = 0.009
g = 9.81

# Datos de entrada (medidos)
y_medido = [0.242, 0.241, 0.239, 0.237, 0.234, 0.232, 0.229, 0.227, 0.224, 0.222]

# Valores de la tabla de Excel para usar en los cálculos (según solicitud)
n_comp_tabla = [0.00939674, 0.00939773, 0.00939973, 0.00940175, 0.00940481, 0.00940688, 0.00941002, 0.00941215,
                0.00941537, 0.00941755]
f_y_tabla = [82.87417, 82.86785, 82.85488, 82.84144, 82.82037, 82.80569, 82.78265, 82.76658, 82.74134, 82.72371]
# Valores de la tabla de Excel para Se (para validar So-Se)
se_tabla = [0.000045, 0.000045, 0.000046, 0.000047, 0.000049, 0.000050, 0.000052, 0.000053, 0.000055, 0.000056]
# Valores de la tabla de Excel para Froude Term (para validar f(y))
froude_tabla = [0.990760, 0.990644, 0.990408, 0.990163, 0.989779, 0.989513, 0.989095, 0.988804, 0.988349, 0.988031]

X_acumulado = 0.0
y_anterior = 0.0
f_y_anterior = 0.0

for i in range(10):
    punto = {}
    y = y_medido[i]

    punto['i'] = i + 1
    punto['y'] = y

    A = b * y
    P = b + 2 * y
    Rh = A / P
    T = b
    V = Q / A

    P_base = b
    P_pared = 2 * y
    P_total = P

    # Cálculo de n_comp para verificación
    n_comp_calc = ((P_base * (n_base ** 1.5) + P_pared * (n_pared ** 1.5)) / P_total) ** (2 / 3)
    # Valor de n_comp a utilizar (de la tabla)
    n_comp_usar = n_comp_tabla[i]

    # Cálculo de Se
    Se_calc = (V * n_comp_usar / (Rh ** (2 / 3))) ** 2
    # Valor de Se a utilizar (de la tabla, para máxima precisión con datos de origen)
    Se_usar = se_tabla[i]

    # Cálculo de Froude Term
    Froude_term_calc = 1 - (Q ** 2 * T) / (g * A ** 3)
    # Valor de Froude a utilizar (de la tabla)
    Froude_usar = froude_tabla[i]

    So_minus_Se = So - Se_usar

    # Valor de f(y) a utilizar (de la tabla)
    f_y_usar = f_y_tabla[i]

    delta_x = None
    X_calc = 0.0

    if i > 0:
        delta_x = (f_y_anterior + f_y_usar) / 2 * (y_anterior - y)
        X_acumulado += delta_x
        X_calc = X_acumulado

    punto['A'] = A
    punto['P'] = P
    punto['Rh'] = Rh
    punto['T'] = T
    punto['V'] = V
    punto['P_base'] = P_base
    punto['P_pared'] = P_pared
    punto['P_total'] = P_total
    punto['n_base'] = n_base
    punto['n_pared'] = n_pared
    punto['n_comp_calc'] = n_comp_calc  # El resultado de la fórmula
    punto['n_comp_usar'] = n_comp_usar  # El valor de la tabla
    punto['Se_calc'] = Se_calc  # El Se calculado
    punto['Se_usar'] = Se_usar  # El Se de la tabla
    punto['Froude_calc'] = Froude_term_calc  # El Froude calculado
    punto['Froude_usar'] = Froude_usar  # El Froude de la tabla
    punto['So_minus_Se'] = So_minus_Se
    punto['f_y_usar'] = f_y_usar
    punto['delta_x'] = delta_x
    punto['X_calc'] = X_calc
    punto['y_anterior'] = y_anterior
    punto['f_y_anterior'] = f_y_anterior

    datos.append(punto)

    # Update for next iteration
    y_anterior = y
    f_y_anterior = f_y_usar


# --- FIN DE DATOS CENTRALIZADOS ---


def add_styled_paragraph(doc, text, size=12, font_name='Times New Roman', line_spacing=1.5, bold=False, indent=None):
    """
    Añade un párrafo al documento con el estilo especificado.
    """
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if indent:
        p.paragraph_format.left_indent = Cm(indent)

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    return p


def set_page_margins(doc):
    """
    Configura los márgenes de página a A4 con las dimensiones especificadas.
    """
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)  # A4
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)


def main():
    """
    Función principal para generar el documento Word.
    """
    doc = docx.Document()

    # Configurar estilo Normal (fuente, tamaño, interlineado)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Configurar márgenes
    set_page_margins(doc)

    # Iterar sobre cada punto y añadir sus cálculos
    for p in datos:
        i = p['i']
        y = p['y']

        # Título del Punto
        add_styled_paragraph(doc, f"Punto {i}: Tramo Y{i} (Tirante y = {y} m)", bold=True)

        # Cálculo de Área (A)
        add_styled_paragraph(doc, f"Cálculo de Área (A{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"A = b * y = {b} m * {y} m", indent=1.5)
        add_styled_paragraph(doc, f"A{i} = {p['A']:.6f} m²", indent=1.5)

        # Cálculo de Perímetro (P)
        add_styled_paragraph(doc, f"Cálculo de Perímetro (P{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"P = b + 2y = {b} m + 2 * {y} m", indent=1.5)
        add_styled_paragraph(doc, f"P{i} = {p['P']:.6f} m", indent=1.5)

        # Cálculo de Radio Hidráulico (Rh)
        add_styled_paragraph(doc, f"Cálculo de Radio Hidráulico (Rh{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"Rh = A / P = {p['A']:.6f} m² / {p['P']:.6f} m", indent=1.5)
        add_styled_paragraph(doc, f"Rh{i} = {p['Rh']:.6f} m", indent=1.5)

        # Cálculo de Espejo de Agua (T)
        add_styled_paragraph(doc, f"Cálculo de Espejo de Agua (T{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"T = b (constante)", indent=1.5)
        add_styled_paragraph(doc, f"T{i} = {p['T']:.6f} m", indent=1.5)

        # Cálculo de Velocidad (V)
        add_styled_paragraph(doc, f"Cálculo de Velocidad (V{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"V = Q / A = {Q} m³/s / {p['A']:.6f} m²", indent=1.5)
        add_styled_paragraph(doc, f"V{i} = {p['V']:.6f} m/s", indent=1.5)

        # Cálculo de Rugosidad Compuesta (n_comp)
        add_styled_paragraph(doc, f"Cálculo de Rugosidad Compuesta (n_comp{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"P_base = b = {p['P_base']:.6f} m", indent=1.5)
        add_styled_paragraph(doc, f"P_pared = 2 * y = 2 * {y} m = {p['P_pared']:.6f} m", indent=1.5)
        add_styled_paragraph(doc, f"P_total = P{i} = {p['P_total']:.6f} m", indent=1.5)
        add_styled_paragraph(doc, f"n_c = [ (P_base * n_base^1.5 + P_pared * n_pared^1.5) / P_total ] ^ (2/3)",
                             indent=1.5)
        add_styled_paragraph(doc,
                             f"n_c = [ ({p['P_base']:.6f} * {p['n_base']}^1.5 + {p['P_pared']:.6f} * {p['n_pared']}^1.5) / {p['P_total']:.6f} ] ^ (2/3)",
                             indent=1.5)
        add_styled_paragraph(doc, f"n_c{i} (calculado) = {p['n_comp_calc']:.10f}", indent=1.5)
        add_styled_paragraph(doc, f"Resultado (Valor de tabla Excel) = {p['n_comp_usar']:.10f}", indent=1.5)

        # Cálculo de Pendiente de Energía (Se)
        add_styled_paragraph(doc, f"Cálculo de Pendiente de Energía (Se{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"Se = (V * n_comp / Rh^(2/3))^2", indent=1.5)
        add_styled_paragraph(doc, f"Se = ({p['V']:.6f} * {p['n_comp_usar']:.10f} / {p['Rh']:.6f}^(2/3))^2", indent=1.5)
        add_styled_paragraph(doc, f"Se{i} (calculado) = {p['Se_calc']:.10f}", indent=1.5)
        add_styled_paragraph(doc, f"Resultado (Valor de tabla Excel) = {p['Se_usar']:.10f}", indent=1.5)

        # Cálculo de Término Froude
        add_styled_paragraph(doc, f"Cálculo de Término Froude (Fr{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"Fr_term = 1 - (Q² * T) / (g * A³)", indent=1.5)
        add_styled_paragraph(doc, f"Fr_term = 1 - ({Q}² * {p['T']:.6f}) / ({g} * {p['A']:.6f}³)", indent=1.5)
        add_styled_paragraph(doc, f"Fr_term{i} (calculado) = {p['Froude_calc']:.10f}", indent=1.5)
        add_styled_paragraph(doc, f"Resultado (Valor de tabla Excel) = {p['Froude_usar']:.10f}", indent=1.5)

        # Cálculo de (So - Se)
        add_styled_paragraph(doc, f"Cálculo de (So - Se){i}:", bold=True, indent=1)
        add_styled_paragraph(doc, f"So - Se = {So} - {p['Se_usar']:.10f}", indent=1.5)
        add_styled_paragraph(doc, f"(So - Se){i} = {p['So_minus_Se']:.10f}", indent=1.5)

        # Cálculo de f(y)
        add_styled_paragraph(doc, f"Cálculo de f(y{i}):", bold=True, indent=1)
        add_styled_paragraph(doc, f"f(y) = (Término Froude) / (So - Se)", indent=1.5)
        add_styled_paragraph(doc, f"f(y{i}) = {p['Froude_usar']:.10f} / {p['So_minus_Se']:.10f}", indent=1.5)
        add_styled_paragraph(doc, f"f(y{i}) (Valor de tabla Excel) = {p['f_y_usar']:.10f}", indent=1.5)

        # Cálculo de Δx (Delta x)
        if p['delta_x'] is not None:
            add_styled_paragraph(doc, f"Cálculo de Δx (Delta x){i}:", bold=True, indent=1)
            add_styled_paragraph(doc, f"Δx = (f(y_anterior) + f(y_actual)) / 2 * (y_anterior - y_actual)", indent=1.5)
            add_styled_paragraph(doc,
                                 f"Δx = ({p['f_y_anterior']:.10f} + {p['f_y_usar']:.10f}) / 2 * ({p['y_anterior']} - {p['y']})",
                                 indent=1.5)
            add_styled_paragraph(doc, f"Δx{i} = {p['delta_x']:.10f} m", indent=1.5)
        else:
            add_styled_paragraph(doc, f"Cálculo de Δx (Delta x){i}:", bold=True, indent=1)
            add_styled_paragraph(doc, "N/A (Punto inicial)", indent=1.5)

        # Cálculo de X (Distancia Acumulada)
        add_styled_paragraph(doc, f"Cálculo de X (Distancia Acumulada){i}:", bold=True, indent=1)
        if p['i'] == 1:
            add_styled_paragraph(doc, f"X{i} = 0.000 m (Punto inicial)", indent=1.5)
        else:
            add_styled_paragraph(doc, f"X_actual = X_anterior + Δx", indent=1.5)
            X_anterior = p['X_calc'] - p['delta_x']
            add_styled_paragraph(doc, f"X{i} = {X_anterior:.10f} m + {p['delta_x']:.10f} m", indent=1.5)
            add_styled_paragraph(doc, f"X{i} = {p['X_calc']:.10f} m", indent=1.5)

        # Espaciador
        add_styled_paragraph(doc, "")

        # Guardar el documento
    file_name = 'memoria_calculo_fgv.docx'
    doc.save(file_name)
    print(f"Documento '{file_name}' generado exitosamente.")


if __name__ == '__main__':
    main()
