# -*- coding: utf-8 -*-
"""
================================================================================
== GENERADOR DE REPORTE: DISEÑO DE MODELO HIDRÁULICO A ESCALA (V5 - DETALLADO) ==
================================================================================

Este script genera una memoria de cálculo ultra detallada que sigue el flujo de
diseño de ingeniería para un canal de máxima eficiencia hidráulica (SME).

La narrativa del reporte sigue el flujo de diseño estándar de ingeniería:
1. Se presentan las especificaciones de diseño del Prototipo: caudal,
   rugosidad y pendiente.
2. Se diseña la sección del Prototipo aplicando el criterio de máxima
   eficiencia hidráulica (b = 2y), calculando y explicando cada parámetro
   hidráulico (y_p, b_p, A_p, V_p, Fr_p) de forma individual.
3. Se presentan todos los parámetros hidráulicos del prototipo ya diseñado.
4. Se definen las restricciones de construcción del Modelo (maqueta).
5. Se calcula la escala de semejanza y se diseñan los parámetros del modelo.
6. Se presenta la validación final de la semejanza dinámica.
"""
import math
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scipy.optimize import fsolve

# --- 1. PARÁMETROS DE ENTRADA DEL ESCENARIO DE DISEÑO ---
# -------------------------------------------------------------------
# --- Parámetros de Diseño del Prototipo (Solo requisitos) ---
PROTOTIPO_INPUTS = {
    "Q_p": 0.87142,  # Caudal de diseño en m³/s
    "n_p": 0.014,  # Rugosidad de Manning (ej. concreto)
    "S": 0.00373,  # Pendiente del canal
    "g": 9.81  # Aceleración de la gravedad
}

# --- Restricciones del Modelo (Maqueta de Laboratorio) ---
MODELO_CONSTRAINTS = {
    "b_m": 0.20,  # Ancho fijo constructivo: 20 cm
}


# -------------------------------------------------------------------

# --- 2. FUNCIONES DE CÁLCULO Y AUXILIARES ---

def manning_sme_for_y(y, Q, n, S):
    """
    Ecuación de Manning para un canal rectangular de máxima eficiencia (b=2y).
    La función está arreglada para ser resuelta con fsolve.
    """
    if y <= 0:
        return float('inf')
    # Para SME rectangular: A = 2y², R = y/2
    return Q - (1 / n) * (2 * y ** 2) * ((y / 2) ** (2 / 3)) * (S ** 0.5)


def add_hyper_detailed_calculation(doc, title, intro_text, steps, final_formula_latex, substitution_latex, values_dict,
                                   result_variable, result_value, result_unit):
    """
    Añade una sección de cálculo con un nivel de detalle extremo al documento.
    """
    doc.add_heading(title, level=3)
    doc.add_paragraph(intro_text)

    p = doc.add_paragraph()
    p.add_run('Derivación Lógica y Algebraica:').bold = True
    for step in steps:
        p_step = doc.add_paragraph(style='List Number')
        parts = step.split('$')
        for i, part in enumerate(parts):
            if i % 2 == 1:
                run = p_step.add_run(part)
                run.font.italic = True
            else:
                p_step.add_run(part)
        p_step.paragraph_format.left_indent = Inches(0.25)

    p = doc.add_paragraph()
    p.add_run('Fórmula Final Aplicable:').bold = True
    p_formula = doc.add_paragraph()
    run = p_formula.add_run(f'${final_formula_latex}$')
    run.font.italic = True
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Sustitución de Datos del Proyecto:').bold = True
    equation_with_values = substitution_latex
    for key, val in values_dict.items():
        str_val = f"{val:.5f}" if isinstance(val, float) else str(val)
        equation_with_values = equation_with_values.replace(key, str_val)
    p_subs = doc.add_paragraph()
    run = p_subs.add_run(f'${equation_with_values}$')
    run.font.italic = True
    p_subs.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Resultado Obtenido:').bold = True
    p_result = doc.add_paragraph()
    p_result.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_result = p_result.add_run(f'{result_variable} = {result_value:.5f} {result_unit}')
    run_result.font.bold = True
    doc.add_paragraph()


# --- 3. GENERADOR DEL INFORME EN WORD ---

def generar_informe_diseno_directo(proto_inputs, modelo_constraints, calculos_completos):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    doc.add_heading('Memoria de Cálculo: Diseño de Modelo de Canal a Escala', level=0)

    # --- Secciones introductorias ---
    doc.add_heading('1. Introducción y Objetivo', level=1)
    doc.add_paragraph(
        "El presente documento detalla la memoria de cálculo para el diseño de un modelo físico a escala reducida "
        "de un tramo de canal rectangular. El objetivo es primero diseñar la sección transversal del canal real (prototipo) "
        "bajo el criterio de máxima eficiencia hidráulica, y luego determinar las dimensiones y condiciones operativas de "
        "una maqueta de laboratorio que lo simule de manera dinámicamente semejante."
    )

    doc.add_heading('2. Datos de Partida y Restricciones', level=1)
    doc.add_heading('2.1. Requisitos de Diseño del Prototipo (Canal Real)', level=2)
    doc.add_paragraph(
        "A continuación, se presentan los parámetros de diseño iniciales para el canal real a simular. Las dimensiones geométricas son desconocidas.")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parámetro'
    hdr_cells[1].text = 'Valor'
    data_rows = [
        ("Caudal de Diseño ($Q_p$)", f"{proto_inputs['Q_p']:.5f} m³/s"),
        ("Rugosidad de Manning ($n_p$)", f"{proto_inputs['n_p']:.4f} (Concreto)"),
        ("Pendiente ($S$)", f"{proto_inputs['S']:.5f} m/m")
    ]
    for key, val in data_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = val

    doc.add_heading('2.2. Restricciones del Modelo (Maqueta de Laboratorio)', level=2)
    table_m = doc.add_table(rows=1, cols=2)
    table_m.style = 'Table Grid'
    hdr_cells_m = table_m.rows[0].cells
    hdr_cells_m[0].text = 'Parámetro'
    hdr_cells_m[1].text = 'Valor'
    row_m = table_m.add_row().cells
    row_m[0].text = "Ancho Fijo del Modelo ($b_m$)"
    row_m[1].text = f"{modelo_constraints['b_m']:.3f} m ({modelo_constraints['b_m'] * 100:.1f} cm)"

    # --- 3. METODOLOGÍA Y DESARROLLO DE CÁLCULOS ---
    doc.add_heading('3. Metodología y Desarrollo de Cálculos', level=1)

    doc.add_heading('3.1. Diseño y Caracterización Detallada del Prototipo (SME)', level=2)

    add_hyper_detailed_calculation(doc, "3.1.1. Cálculo del Tirante Óptimo ($y_p$)",
                                   "El primer paso es diseñar la sección transversal del canal. Se utiliza el criterio de máxima eficiencia hidráulica, que minimiza el perímetro mojado para un área de flujo dada, resultando en la máxima velocidad y el menor costo de construcción.",
                                   [
                                       "Para un canal rectangular, la condición de máxima eficiencia se cumple cuando el ancho es el doble del tirante: $b = 2y$.",
                                       "Esta condición simplifica el Área ($A = 2y^2$) y el Radio Hidráulico ($R_h = y/2$).",
                                       "Sustituyendo estas relaciones en la ecuación de Manning, se obtiene una ecuación con una única incógnita, el tirante $y_p$."
                                   ],
                                   "Q_p = \\frac{1}{n_p} (2y_p^2) \\left( \\frac{y_p}{2} \\right)^{2/3} S^{1/2}",
                                   "Qp_val = \\frac{1}{np_val} (2y_p^2) \\left( \\frac{y_p}{2} \\right)^{2/3} S_val^{1/2}",
                                   {'Qp_val': proto_inputs['Q_p'], 'np_val': proto_inputs['n_p'],
                                    'S_val': proto_inputs['S']},
                                   "y_p", calculos_completos['prototipo']['y_p'], "m"
                                   )

    add_hyper_detailed_calculation(doc, "3.1.2. Cálculo del Ancho de Solera Óptimo ($b_p$)",
                                   "El ancho de solera se deriva directamente de la condición de máxima eficiencia hidráulica para una sección rectangular.",
                                   ["La condición establece que el ancho debe ser el doble del tirante para minimizar el perímetro mojado."],
                                   "b_p = 2 \\cdot y_p",
                                   "b_p = 2 \\cdot yp_val",
                                   {'yp_val': calculos_completos['prototipo']['y_p']},
                                   "b_p", calculos_completos['prototipo']['b_p'], "m"
                                   )

    add_hyper_detailed_calculation(doc, "3.1.3. Cálculo del Área Hidráulica ($A_p$)",
                                   "Corresponde al área de la sección transversal del flujo, y se calcula con las dimensiones ya diseñadas del canal.",
                                   ["El área de una sección rectangular es el producto de su base (ancho de solera) y su altura (tirante)."],
                                   "A_p = b_p \\cdot y_p",
                                   "A_p = bp_val \\cdot yp_val",
                                   {'bp_val': calculos_completos['prototipo']['b_p'],
                                    'yp_val': calculos_completos['prototipo']['y_p']},
                                   "A_p", calculos_completos['prototipo']['A'], "m²"
                                   )

    add_hyper_detailed_calculation(doc, "3.1.4. Cálculo de la Velocidad del Flujo ($V_p$)",
                                   "La velocidad media del flujo se determina a partir de la ecuación de continuidad.",
                                   ["La continuidad establece que el caudal ($Q$) es igual al producto del área de flujo ($A$) y la velocidad media ($V$)."],
                                   "V_p = \\frac{Q_p}{A_p}",
                                   "V_p = \\frac{Qp_val}{Ap_val}",
                                   {'Qp_val': calculos_completos['prototipo']['Q_p'],
                                    'Ap_val': calculos_completos['prototipo']['A']},
                                   "V_p", calculos_completos['prototipo']['V'], "m/s"
                                   )

    add_hyper_detailed_calculation(doc, "3.1.5. Cálculo del Número de Froude ($Fr_p$)",
                                   "Se calcula el número de Froude para caracterizar el régimen de flujo del prototipo. Este número es fundamental para la semejanza dinámica.",
                                   [
                                       "El número de Froude ($Fr$) es la relación entre las fuerzas de inercia (representadas por la velocidad $V$) y las fuerzas de gravedad (representadas por la celeridad de una onda $\\sqrt{g \\cdot y}$).",
                                       "Un valor $Fr < 1$ indica régimen subcrítico, $Fr = 1$ es crítico, y $Fr > 1$ es supercrítico."
                                   ],
                                   "Fr_p = \\frac{V_p}{\\sqrt{g \\cdot y_p}}",
                                   "Fr_p = \\frac{Vp_val}{\\sqrt{g_val \\cdot yp_val}}",
                                   {'Vp_val': calculos_completos['prototipo']['V'],
                                    'g_val': calculos_completos['prototipo']['g'],
                                    'yp_val': calculos_completos['prototipo']['y_p']},
                                   "Fr_p", calculos_completos['prototipo']['Fr'], "(adimensional)"
                                   )

    doc.add_heading('3.2. Resumen de Parámetros del Prototipo Diseñado', level=2)
    doc.add_paragraph("Concluido el diseño, las propiedades hidráulicas completas del canal real son:")
    table_p_full = doc.add_table(rows=1, cols=2)
    table_p_full.style = 'Table Grid'
    # ... (El resto de la función se mantiene y renumera)
    hdr_cells_p_full = table_p_full.rows[0].cells
    hdr_cells_p_full[0].text = 'Parámetro Calculado'
    hdr_cells_p_full[1].text = 'Valor'
    proto_results = calculos_completos['prototipo']
    data_rows_p_full = [
        ("Tirante Óptimo ($y_p$)", f"{proto_results['y_p']:.4f} m"),
        ("Ancho de Solera Óptimo ($b_p$)", f"{proto_results['b_p']:.4f} m"),
        ("Área Hidráulica ($A_p$)", f"{proto_results['A']:.4f} m²"),
        ("Velocidad del Flujo ($V_p$)", f"{proto_results['V']:.4f} m/s"),
        ("Número de Froude ($Fr_p$)", f"{proto_results['Fr']:.4f}")
    ]
    for key, val in data_rows_p_full:
        row_cells = table_p_full.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = val

    doc.add_heading('3.3. Determinación de Escalas de Semejanza', level=2)

    add_hyper_detailed_calculation(doc, "3.3.1. Escala de Longitud ($\\lambda_L$)",
                                   "Define la relación geométrica entre el prototipo diseñado y el modelo. Se calcula a partir de las dimensiones conocidas.",
                                   [
                                       "La escala de longitud ($\\lambda_L$) es el cociente entre una dimensión homóloga del prototipo y el modelo.",
                                       "Se utiliza el ancho de la solera como dimensión característica para establecer esta relación."
                                   ],
                                   "\\lambda_L = \\frac{b_p}{b_m}",
                                   "\\lambda_L = \\frac{bp_val}{bm_val}",
                                   {'bp_val': calculos_completos['prototipo']['b_p'],
                                    'bm_val': modelo_constraints['b_m']},
                                   "\\lambda_L", calculos_completos['escala']['lambda_L'], "(adimensional)"
                                   )

    doc.add_heading('3.4. Diseño de Parámetros del Modelo a Escala', level=2)
    doc.add_paragraph(
        "Con la escala de longitud definida, se procede a calcular los parámetros requeridos para la construcción y operación de la maqueta.")

    add_hyper_detailed_calculation(doc, "3.4.1. Tirante del Modelo ($y_m$)",
                                   "Determina la profundidad del agua que se debe mantener en la maqueta para simular el tirante del prototipo.",
                                   ["El tirante del modelo se obtiene escalando geométricamente el tirante del prototipo usando la escala de longitud."],
                                   "y_m = \\frac{y_p}{\\lambda_L}",
                                   "y_m = \\frac{yp_val}{lambda_L_val}",
                                   {'yp_val': calculos_completos['prototipo']['y_p'],
                                    'lambda_L_val': calculos_completos['escala']['lambda_L']},
                                   "y_m", calculos_completos['modelo']['y'], "m"
                                   )

    add_hyper_detailed_calculation(doc, "3.4.2. Rugosidad Requerida del Modelo ($n_m$)",
                                   "Determina el coeficiente de Manning que debe tener el material de la maqueta para garantizar la semejanza dinámica.",
                                   [
                                       "Para que las leyes de Manning y Froude se cumplan simultáneamente, la rugosidad debe ser escalada.",
                                       "La relación de escalas de rugosidad es $n_r = \\frac{n_p}{n_m} = \\lambda_L^{1/6}$.",
                                       "Despejando $n_m$ se obtiene la rugosidad que debe tener el material del modelo."
                                   ],
                                   "n_m = \\frac{n_p}{\\lambda_L^{1/6}}",
                                   "n_m = \\frac{np_val}{lambda_L_val^{1/6}}",
                                   {'np_val': calculos_completos['prototipo']['n_p'],
                                    'lambda_L_val': calculos_completos['escala']['lambda_L']},
                                   "n_m", calculos_completos['modelo']['n'], "(adimensional)"
                                   )

    add_hyper_detailed_calculation(doc, "3.4.3. Caudal Operativo del Modelo ($Q_m$)",
                                   "Calcula el caudal que se debe suministrar a la maqueta para simular las condiciones del prototipo.",
                                   [
                                       "La escala de caudal se deriva de las escalas de longitud y velocidad: $\\lambda_Q = \\lambda_A \\cdot \\lambda_V$.",
                                       "Sabiendo que $\\lambda_A = \\lambda_L^2$ y $\\lambda_V = \\lambda_L^{1/2}$, se obtiene $\\lambda_Q = \\lambda_L^{2.5}$.",
                                       "El caudal del modelo se calcula escalando inversamente el caudal del prototipo."
                                   ],
                                   "Q_m = \\frac{Q_p}{\\lambda_L^{2.5}}",
                                   "Q_m = \\frac{Qp_val}{lambda_L_val^{2.5}}",
                                   {'Qp_val': calculos_completos['prototipo']['Q_p'],
                                    'lambda_L_val': calculos_completos['escala']['lambda_L']},
                                   "Q_m", calculos_completos['modelo']['Q'], "m³/s"
                                   )

    # --- 4. Resultados y Validación ---
    doc.add_heading('4. Resumen de Parámetros de Diseño del Modelo', level=1)
    doc.add_paragraph(
        "La tabla resume las características finales que deberá tener la maqueta de laboratorio para simular correctamente el prototipo.")
    table_p = doc.add_table(rows=1, cols=2)
    table_p.style = 'Table Grid'
    hdr_cells_p = table_p.rows[0].cells
    hdr_cells_p[0].text = 'Parámetro del Modelo'
    hdr_cells_p[1].text = 'Valor de Diseño'
    modelo_results = calculos_completos['modelo']
    data_rows_p = [
        ("Ancho de Solera ($b_m$)", f"{modelo_results['b']:.4f} m"),
        ("Tirante Normal ($y_m$)", f"{modelo_results['y']:.4f} m ({modelo_results['y'] * 100:.2f} cm)"),
        ("Caudal de Operación ($Q_m$)", f"{modelo_results['Q']:.5f} m³/s ({modelo_results['Q'] * 1000:.2f} L/s)"),
        ("Rugosidad Requerida ($n_m$)", f"{modelo_results['n']:.5f} (compatible con PVC/acrílico)"),
        ("Velocidad del Flujo ($V_m$)", f"{modelo_results['V']:.4f} m/s"),
        ("Número de Froude ($Fr_m$)", f"{modelo_results['Fr']:.4f}")
    ]
    for key, val in data_rows_p:
        row_cells = table_p.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = val

    doc.add_heading('5. Verificación de Semejanza Dinámica y Conclusión', level=1)
    doc.add_paragraph(
        "El paso final es verificar que la semejanza de Froude, base de todo el escalamiento, se cumple. Se calculan los números de Froude para ambos sistemas de forma independiente y se comparan.")

    p = doc.add_paragraph()
    p.add_run(f"Número de Froude del Prototipo (Fr_p): ").bold = True
    p.add_run(f"{calculos_completos['prototipo']['Fr']:.5f}")

    p = doc.add_paragraph()
    p.add_run(f"Número de Froude del Modelo (Fr_m): ").bold = True
    p.add_run(f"{calculos_completos['modelo']['Fr']:.5f}")

    doc.add_paragraph()
    if math.isclose(calculos_completos['prototipo']['Fr'], calculos_completos['modelo']['Fr'], rel_tol=1e-5):
        p_val = doc.add_paragraph()
        p_val.add_run(
            "VALIDACIÓN EXITOSA: Los números de Froude son idénticos. El diseño del modelo a escala es consistente y representará fielmente la dinámica del prototipo.").bold = True
    else:
        p_val = doc.add_paragraph()
        p_val.add_run(
            "ERROR DE VALIDACIÓN: Los números de Froude no coinciden, indicando un error en la lógica de cálculo.").bold = True

    filename = 'Reporte_Diseno_SME_Modelo_a_Escala.docx'
    doc.save(filename)
    return filename


# --- 4. FLUJO PRINCIPAL DE EJECUCIÓN ---
if __name__ == "__main__":
    print("Iniciando diseño de modelo a escala y generación de memoria de cálculo...")

    # --- Realizar todos los cálculos necesarios ---
    calculos = {"prototipo": {}, "modelo": {}, "escala": {}}

    # A. Diseñar el Prototipo para Máxima Eficiencia Hidráulica
    print("Paso 1: Diseñando sección de máxima eficiencia para el prototipo...")
    y_p_sme = fsolve(manning_sme_for_y, 0.5,  # Estimación inicial de 50cm
                     args=(PROTOTIPO_INPUTS['Q_p'], PROTOTIPO_INPUTS['n_p'],
                           PROTOTIPO_INPUTS['S']))[0]

    b_p_sme = 2 * y_p_sme

    proto_full_params = {**PROTOTIPO_INPUTS, "y_p": y_p_sme, "b_p": b_p_sme}

    proto_A = proto_full_params['b_p'] * proto_full_params['y_p']
    proto_V = proto_full_params['Q_p'] / proto_A
    proto_Fr = proto_V / math.sqrt(proto_full_params['g'] * proto_full_params['y_p'])
    calculos["prototipo"] = {**proto_full_params, "A": proto_A, "V": proto_V, "Fr": proto_Fr}

    print(f"   -> Tirante Óptimo (y_p) calculado: {y_p_sme:.4f} m")
    print(f"   -> Ancho Óptimo (b_p) calculado: {b_p_sme:.4f} m")

    # B. Calcular Escalas
    print("Paso 2: Calculando escalas de semejanza...")
    lambda_L = calculos['prototipo']['b_p'] / MODELO_CONSTRAINTS['b_m']
    calculos["escala"] = {"lambda_L": lambda_L}
    print(f"   -> Escala de Longitud (λ_L): {lambda_L:.2f}")

    # C. Calcular Parámetros del Modelo
    print("Paso 3: Diseñando parámetros del modelo a escala...")
    modelo_b = MODELO_CONSTRAINTS['b_m']
    modelo_y = calculos['prototipo']['y_p'] / lambda_L
    modelo_n = calculos['prototipo']['n_p'] / (lambda_L ** (1 / 6))
    modelo_Q = calculos['prototipo']['Q_p'] / (lambda_L ** 2.5)

    modelo_A = modelo_b * modelo_y
    modelo_V = modelo_Q / modelo_A
    modelo_Fr = modelo_V / math.sqrt(calculos['prototipo']['g'] * modelo_y)

    calculos["modelo"] = {
        "b": modelo_b, "y": modelo_y, "n": modelo_n, "Q": modelo_Q,
        "A": modelo_A, "V": modelo_V, "Fr": modelo_Fr
    }
    print("   -> Parámetros del modelo calculados exitosamente.")

    print("Paso 4: Generando informe en Word...")

    try:
        nombre_archivo = generar_informe_diseno_directo(PROTOTIPO_INPUTS, MODELO_CONSTRAINTS, calculos)
        print(f"\n¡Éxito! Se ha generado el informe: '{nombre_archivo}'")
        print("El archivo detalla el proceso de diseño desde el canal real hasta la maqueta de laboratorio.")
        # Abrir el archivo automáticamente (opcional)
        try:
            os.startfile(nombre_archivo)
        except AttributeError:
            print(f"Para ver el informe, abre el archivo '{nombre_archivo}' manualmente.")
    except Exception as e:
        print(f"\nError: No se pudo generar el archivo de Word. Causa: {e}")

