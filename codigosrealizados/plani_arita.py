import os
import logging
import json
import re
from pathlib import Path
from datetime import datetime

# Importación condicional para manejo de errores (Ingeniería Pesimista)
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("CRITICAL ERROR: La librería 'python-docx' no está instalada.")
    print("Por favor ejecuta: pip install python-docx")
    exit(1)

# --- CONFIGURACIÓN DE LOGGING ---
log_filename = f"wbs_generation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_filename, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class WBSReportGenerator:
    """
    Generador de Reportes EDT Inteligente.
    Parsea jerarquías de presupuestos (S10/Excel) para generar Diccionarios EDT completos en Word.
    """

    def __init__(self, output_dir="resultados_edt"):
        self.output_dir = Path(output_dir)
        self._ensure_directories()

        # Base de conocimiento ampliada para tu proyecto específico
        self.knowledge_base = {
            "OBRAS PROVISIONALES": "Comprende todas las construcciones e instalaciones temporales necesarias para el servicio de la obra, incluyendo cartel, almacén, guardianía y servicios higiénicos.",
            "DESMONTAJES": "Trabajos de retiro de estructuras existentes (arcos, postes, coberturas) que interfieren con el nuevo proyecto, recuperando materiales útiles y eliminando los desechos.",
            "DEMOLICIONES": "Demolición de estructuras de concreto y albañilería existentes. Incluye la fragmentación y acondicionamiento del material para su eliminación.",
            "SEGURIDAD Y SALUD": "Implementación de los estándares de seguridad (Norma G.050), incluyendo EPPs, señalización, capacitaciones y plan de contingencia COVID-19/Emergencias.",
            "MOVIMIENTO DE TIERRAS": "Excavaciones, cortes, rellenos y nivelación del terreno necesarios para alcanzar las cotas de fundación y subrasante indicadas en los planos.",
            "CONCRETO SIMPLE": "Elementos de concreto sin refuerzo estructural (solados, falsos pisos, dados de anclaje) destinados a recibir cargas menores o mejorar el terreno.",
            "CONCRETO ARMADO": "Estructuras principales (zapatas, columnas, vigas, losas, placas) de concreto f'c=210 kg/cm2 reforzado con acero corrugado grado 60.",
            "COBERTURA": "Suministro e instalación de techos, incluyendo estructuras metálicas de soporte y coberturas (calamina, autoportante, lona PVC) según diseño.",
            "PINTURA": "Aplicación de pintura en muros, cielorrasos, carpintería metálica y demarcación de losa deportiva, garantizando protección y estética.",
            "CARPINTERIA METALICA": "Fabricación e instalación de elementos de acero como arcos, postes de vóley, barandas, puertas y protecciones metálicas.",
            "GRASS SINTETICO": "Instalación de sistema de césped artificial para campo deportivo, incluyendo preparación de base, pegado de uniones y líneas de demarcación.",
            "INSTALACIONES ELECTRICAS": "Redes de alimentación, tableros, iluminación (reflectores LED, luminarias) y sistema de puesta a tierra para garantizar operatividad nocturna.",
            "APARATOS SANITARIOS": "Suministro e instalación de inodoros, lavatorios, urinarios y grifería en los módulos de servicios higiénicos.",
            "SISTEMA DE DESAGUE": "Red de recolección de aguas servidas y pluviales, incluyendo tuberías PVC, cajas de registro y conexión a la red pública.",
            "JUEGOS INFANTILES": "Suministro, ensamblaje y fijación segura de módulos recreativos para niños, cumpliendo normas de seguridad y ergonomía.",
            "EQUIPAMIENTO URBANO": "Instalación de bancas, basureros, pérgolas y elementos de paisajismo que complementan la losa deportiva.",
            "AREAS VERDES": "Preparación de terreno agrícola, sembrado de césped natural, árboles y plantas ornamentales para mejoramiento ambiental.",
            "VARIOS": "Elementos complementarios como juntas de dilatación, placas recordatorias y acabados finales no incluidos en otras partidas."
        }

    def _ensure_directories(self):
        try:
            if not self.output_dir.exists():
                self.output_dir.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            logger.critical(f"Error fatal al crear directorios: {e}")
            raise

    def parse_complex_input(self):
        """
        Algoritmo avanzado de parsing.
        Lee el texto crudo y asocia las actividades (líneas sin 'x')
        al último paquete de trabajo (línea con 'x') detectado.
        """
        raw_text = """
x											OBRAS PROVISIONALES Y TRABAJOS PRELIMINARES
	4	1	0	0	0	0				C	CARTEL DE IDENTIFICACION DE LA OBRA DE 3.60 X 2.40 m
	4	1	0	0	0	0				C	ALQUILER DE ALMACEN, GUARDIANIA Y OFICINA
	4	1	0	0	0	0				C	MOVILIZACION Y DESMOVILIZACION DE EQUIPOS Y MAQUINARIAS
	4	1	0	0	0	0				C	FLETE TERRESTRE
	4	1	0	0	0	0				C	SERVICIOS HIGIENICOS PARA LA OBRA
	4	1	0	0	0	0				C	FLUIDO ELECTRICO PARA LA OBRA
	4	1	0	0	0	0				C	TRAZO, REPLANTEO PRELIMINAR Y CONTROL TOPOGRAFICO DE OBRA
	4	1	0	0	0	0				C	CERCO PERMIETRICO PROVISIONAL
x											DESMONTAJES 
	4	1	1	0	0	0				C	DESMONTAJE DE ARCO METALICO EXISTENTE
	4	1	1	0	0	0				C	DESMONTAJE DE POSTES DE CONCRETO
	4	1	1	0	0	0				C	DESMONTAJE DE COBERTURA DE GRADAS
	4	1	1	0	0	0				C	DESMONTAJE DE CERCO METALICO EXISTENTE
	4	1	1	0	0	0				C	DESMONTAJE DE PORTON METALICO DE INGRESO
	4	1	1	0	0	0				C	DESMONTAJE DE PUERTAS
	4	1	1	0	0	0				C	DESMONTAJE DE VENTANAS
	4	1	1	0	0	0				C	DESMONTAJE DE APARATOS SANITARIOS
x											DEMOLICIONES
	4	1	2	0	0	0				C	DEMOLICION DE CIMIENTOS
	4	1	2	0	0	0				C	DEMOLICION DE GRADERIAS DE CONCRETO
	4	1	2	0	0	0				C	DEMOLICION DE LOSA DEPORTIVA
	4	1	2	0	0	0				C	DEMOLICION DE VEREDAS DE 0.10 m
	4	1	2	0	0	0				C	DEMOLICION DE MUROS DE LADRILLO
	4	1	2	0	0	0				C	DEMOLICION DE LOSA ALIGERADA H=0.20M
	4	1	2	0	0	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
x											SEGURIDAD Y SALUD OCUPACIONAL
	4	2	0	0	0	0				C	ELABORACION E IMPLEMENTACION DE PLAN DE SEGURIDAD
	4	2	0	0	0	0				C	EQUIPOS E IMPLEMENTOS PERSONALES DE SEGURIDAD
	4	2	0	0	0	0				C	EQUIPOS DE PROTECCION COLECTIVA
	4	2	0	0	0	0				C	CAPACITACION EN SEGURIDAD Y SALUD
x											COBERTURA AUTOPORTANTE
										C	ESTRUCTURAS
x										C	MOVIMIENTO DE TIERRAS
	4	3	3	1	1	0				C	EXCAVACION MANUAL PARA ZAPATAS EN TERRENO NORMAL
	4	3	3	1	1	0				C	RELLENO COMPACTADO CON MATERIAL PROPIO
	4	3	3	1	1	0				C	AFIRMADO COMPACTADO DE 8'' 
	4	3	3	1	1	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
x										C	CONCRETO SIMPLE
	4	3	3	1	2	0				C	SOLADOS DE CONCRETO C:H 1:12, E=4" PARA ZAPATAS
x										C	CONCRETO ARMADO
										C	ZAPATAS
	4	3	3	1	3	0				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	3	3	1	3	0				C	CONCRETO f'c=210 kg/cm2 EN ZAPATAS
										C	COLUMNAS
	4	3	3	1	3	2				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	3	3	1	3	2				C	ENCOFRADO Y DESENCOFRADO CARAVISTA EN COLUMNAS
	4	4	6	1	3	2				C	CONCRETO f'c=210 kg/cm2 EN COLUMNAS
	4	3	3	1	3	2				C	CURADO DE CONCRETO EN COLUMNAS
										C	VIGAS
	4	3	3	1	3	3				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	3	3	1	3	3				C	ENCOFRADO Y DESENCOFRADO CARAVISTA EN VIGAS
	4	3	3	1	3	3				C	CONCRETO f'c=210 kg/cm2 EN VIGAS
	4	3	3	1	3	3				C	CURADO DE CONCRETO EN VIGAS
x										C	COBERTURA METALICA AUTOPORTANTE
	4	3	3	1	3	4				C	SUMINISTRO E INSTALACION DE COBERTURA METALICA AUTOPORTANTE
	4	3	3	1	3	4				C	SUMINISTRO E INSTALACION DE MUROS TIMPANOS
										C	ARQUITECTURA Y ACABADOS
x										C	PINTURA
	4	3	3	4	1	0				C	PINTURA LATEX 2 MANOS EN COLUMNAS Y VIGAS
										C	LOSA DEPORTIVA DE GRASS SINTETICO
										C	ESTRUCTURAS
x										C	MOVIMIENTO DE TIERRAS
	4	3	4	1	1	0				C	EXCAVACION MANUAL PARA DADOS DE CONCRETO
	4	3	4	1	1	0				C	CONFORMACION Y COMPACTACION DE SUB RASANTE, CON EQUIPO
	4	3	4	1	1	0				C	BASE GRANULAR COMPACTADA H=0.15m CON EQUIPO
x										C	CONCRETO SIMPLE
	4	3	4	1	2	0				C	DADOS PARA POSTES DE ACERO PARA MALLA DE NYLON: CONCRETO f'c=175 kg/cm2
	4	3	4	1	2	0				C	DADOS PARA POSTES DE ACERO PARA MALLA DE VOLEY: CONCRETO f'c=175 kg/cm2
										C	ARQUITECTURA Y ACABADOS
x										C	CARPINTERIA METALICA
	4	3	4	4	2	0				C	PERNO DE ANCLAJE DE Ø 5/8" PARA DADOS DE CONCRETO PARA POSTES DE ACERO MALLA NYLON
	4	3	4	4	2	0				C	TUBOS DE ANCLAJE PARA PARANTES DE VOLEY
	4	3	4	4	2	0				C	TAPA CIRCULAR EN TUBO DE ANCLAJE PARA PARANTES DE VOLEY
	4	3	4	4	2	0				C	ARGOLLAS DE FIERRO PARA TEMPLADO DE PARANTES DE VOLEY
	4	3	4	4	2	0				C	BARRA DE F°G° PARA INGRESO A CANCHA DE GRASS SINTETICO
	4	3	4	4	2	0				C	SUMINISTRO E INSTALACION DE POSTES DE ACERO PARA MALLA DE NYLON
x										C	PINTURA
	4	3	4	4	1	0				C	PINTURA ESMALTE 2 MANOS EN CARPINTERIA METALICA
	4	3	4	4	1	0				C	PINTURA ESMALTE 2 MANOS P/DEMARCACION EN GRASS SINTETICO
x										C	GRASS SINTETICO
	4	3	4	4	3	0				C	LIMPIEZA Y PREPARACION PARA GRASS SINTETICO
	4	3	4	4	3	0				C	SUMINISTRO E INSTALACION DE GRASS SINTETICO EN LOSA DEPORTIVA
x										C	VARIOS
	4	3	4	4	4	0				C	SUMINISTRO E INSTALACION DE ARCOS DE FULBITO (INCLUYE MALLA)
	4	3	4	4	4	0				C	SUMINISTRO E INSTALACION DE POSTES DE VOLEY (INCLUYE MALLA)
										C	SUMINISTRO E INSTALACION DE MALLA DE NYLON ALQUITRANADA N° 96
										C	TRIBUNAS DE CONCRETO
										C	ESTRUCTURAS
x										C	MOVIMIENTO DE TIERRAS
	4	3	5	1	1	0				C	EXCAVACION MANUAL PARA CIMIENTOS EN TERRENO NORMAL
	4	3	3	1	1	0				C	RELLENO COMPACTADO CON MATERIAL PROPIO
	4	3	5	1	1	0				C	RELLENO COMPACTADO CON AFIRMADO
	4	3	5	1	1	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
x										C	CONCRETO SIMPLE
	4	3	5	1	2	0				C	CIMIENTOS CORRIDOS MEZCLA 1:10 CEMENTO:HORMIGON 30% PIEDRA
x										C	CONCRETO ARMADO
										C	PLACAS
	4	3	3	1	3	3				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	3	5	1	3	1				C	ENCOFRADO Y DESENCOFRADO CARAVISTA EN PLACAS DE TRIBUNA
	4	3	5	1	3	1				C	CONCRETO f'c=210 kg/cm2 EN TRIBUNAS
	4	3	5	1	3	1				C	CURADO DE CONCRETO EN TRIBUNAS
										C	LOSAS MACIZAS
	4	3	5	1	3	1				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	3	5	1	3	2				C	ENCOFRADO Y DESENCOFRADO EN LOSAS MACIZAS
	4	3	5	1	3	2				C	CONCRETO f'c=210 kg/cm2 EN LOSAS MACIZAS
	4	3	5	1	3	2				C	CURADO DE CONCRETO EN LOSAS MACIZAS
	4	3	5	1	3	2				C	ACABADO SEMIPULIDO EN LOSAS MACIZAS C/MORTERNO 1:2, E=1cm
										C	ARQUITECTURA Y ACABADOS
	4	3	5	4	0	0				C	JUNTA DE DILATACION CON ESPUMA PLASTICA + JEBE MICROPOROSO
											S2
											SERVICIOS HIGIENICOS Y MODULO ADMINISTRATIVO
										C	ESTRUCTURAS
x										C	MOVIMIENTO DE TIERRAS
	4	4	6	1	1	0				C	EXCAVACION MANUAL PARA CIMIENTOS EN TERRENO NORMAL
	4	4	6	1	1	0				C	RELLENO COMPACTADO CON MATERIAL PROPIO
	4	4	6	1	1	0				C	AFIRMADO COMPACTADO DE 8'' 
	4	4	6	1	1	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
x										C	CONCRETO SIMPLE
	4	4	6	1	2	0				C	CIMIENTOS CORRIDOS MEZCLA 1:10 CEMENTO:HORMIGON 30% PIEDRA
	4	4	6	1	2	0				C	FALSO PISO MEZCLA DE CONCRETO C:H, 1:8, E=4"
	4	4	6	1	2	0				C	CONCRETO f'c=175 kg/cm2 EN PISO DE CISTERNA
	4	4	6	1	2	0				C	ENCOFRADO Y DESENCOFRADO DE SOBRECIMIENTOS
	4	4	6	1	2	0				C	CONCRETO f'c=140 kg/cm2 + 25% P.M. EN SOBRECIMIENTOS
x										C	CONCRETO ARMADO
										C	ZAPATAS
	4	4	6	1	3	1				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	1				C	CONCRETO f'c=210 kg/cm2 EN ZAPATAS
										C	COLUMNAS
	4	4	6	1	3	2				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	2				C	ENCOFRADO Y DESENCOFRADO DE COLUMNAS
	4	4	6	1	3	2				C	CONCRETO f'c=210 kg/cm2 EN COLUMNAS
										C	COLUMNETAS DE AMARRE
	4	4	6	1	3	3				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	3				C	ENCOFRADO Y DESENCOFRADO DE COLUMETAS DE AMARRE
	4	4	6	1	3	3				C	CONCRETO f'c=175 kg/cm2 EN COLUMNETAS DE AMARRE
										C	VIGAS
	4	4	6	1	3	4				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	4				C	ENCOFRADO Y DESENCOFRADO DE VIGAS
	4	4	6	1	3	4				C	CONCRETO f'c=210 kg/cm2 EN VIGAS
										C	VIGUETAS DE AMARRE
	4	4	6	1	3	5				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	5				C	ENCOFRADO Y DESENCOFRADO EN VIGUETAS DE AMARRE
	4	4	6	1	3	5				C	CONCRETO f'c=175 kg/cm2 EN VIGUETAS DE AMARRE
										C	LOSAS ALIGERADAS
	4	4	6	1	3	6				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	6				C	ENCOFRADO Y DESENCOFRADO EN LOSAS ALIGERADAS
	4	4	6	1	3	6				C	CONCRETO f'c=210 kg/cm2 EN LOSAS ALIGERADAS
	4	4	6	1	3	6				C	LADRILLO HUECO DE ARCILLA 15x30x30cm PARA LOSA ALIGERADA
										C	LOSA MACIZA EN CISTERNA
	4	4	6	1	3	7				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	7				C	ENCOFRADO Y DESENCOFRADO EN LOSAS MACIZAS
	4	4	6	1	3	7				C	CONCRETO f'c=210 kg/cm2 EN LOSAS MACIZAS
	4	4	6	1	3	7				C	MURO DE LADRILLO PARA CISTERNA
										C	LOSA PARA LAVATORIOS
	4	4	6	1	3	8				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	8				C	ENCOFRADO Y DESENCOFRADO EN LAVATORIO
	4	4	6	1	3	8				C	CONCRETO f'c=210 kg/cm2 EN LAVATORIO
										C	BANCA PARA VESTIDORES
	4	4	6	1	3	9				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	9				C	ENCOFRADO Y DESENCOFRADO EN BANCA DE VESTIDORES
	4	4	6	1	3	9				C	CONCRETO f'c=210 kg/cm2 EN BANCA PARA VESTIDORES
										C	CASETA DE ELECTROBOMBA
	4	4	6	1	3	10				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	4	6	1	3	10				C	ENCOFRADO Y DESENCOFRADO
	4	4	6	1	3	10				C	CONCRETO f'c=210 kg/cm2 EN CASETA DE ELECTROBOMBA
	4	4	6	1	3	10				C	ACABADO SEMIPULIDO EN PISOS Y LOSA C/MORTERO 1:2, E=1cm
										C	ARQUITECTURA Y ACABADOS
x										C	MUROS Y TABIQUES
	4	4	6	4	1	0				C	MURO DE LADRILLO KING KONG 18 HUECOS DE SOGA C:C:A - 1:1:4 x 1.5cm
	4	4	6	4	1	0				C	COBERTURA DE TECHO CON LADRILLO PASTELERO 24cm x 24cm x 3cm
x										C	REVOQUES Y ENLUCIDOS
	4	4	6	4	2	0				C	TARRAJEO EN MUROS INTERIORES Y EXTERIORES CON C:A - 1:5
	4	4	6	4	2	0				C	VESTIDURA DE DERRAMES
										C	FORJADO Y REVESTIMIENTO DE MESA C:A, 1:4
x										C	CIELORASOS
	4	4	6	4	3	0				C	CIELO RASOS CON MEZCLA DE CEMENTO:ARENA, 1:5
x										C	PISOS
	4	4	6	4	4	0				C	CONTRAPISOS e=40mm, C:A=1:4
										C	ZOCALOS Y CONTRAZOCALOS
	4	4	6	4	5	0				C	ZOCALO DE CERAMICA 0.45x0.45m, H=1.50m
	4	4	6	4	5	0				C	PISO DE CERAMICA 0.45x0.45m ALTO TRANSITO
	4	4	6	4	5	0				C	PISO DE PORCELANATO 0.60x0.60m ALTO TRANSITO
	4	4	6	4	5	0				C	CONTRAZOCALOS DE CEMENTO PULIDO, C:A=1:3, e=1.5cm, h=30cm
x										C	CARPINTERIA METALICA
	4	4	6	4	6	0				C	MODULO METALICO PARA INODOROS SS.HH. MUJERES (SEGÚN DISEÑO)
	4	4	6	4	6	0				C	MODULO METALICO PARA INODOROS SS.HH. HOMBRES (SEGÚN DISEÑO)
x										C	CARPINTERIA DE ALUMINIO
	4	4	6	4	7	0				C	SUMINISTRO E INSTALACION DE VENTANAS DE CRISTAL TEMPLADO 6mm. C/ACCES. METALICO, SIST. CORREDIZO
	4	4	6	4	7	0				C	CANTONERA DE ALUMINIO RANURADO PARA DESNIVEL EN PISOS INTERIORES
x										C	CARPINTERIA DE MADERA
	4	4	6	4	8	0				C	PUERTA CONTRAPLACADA CON TRIPLAY DE 4MM,INCLUYE MARCO Y REJILLA METALICA DE 0.35 x 0.40
	4	4	6	4	8	0				C	PUERTA DE MADERA CEDRO MACHIMBRADA E=2" C/MARCO (UNA HOJA)
x										C	CERRAJERIA
	4	4	6	4	9	0				C	SUMINISTRO E INSTALACION DE CERRADURA DE ACERO INOXIDABLE 3 GOLPES 
	4	4	6	4	9	0				C	CERRADURA DE EMBUTIR SIMPLE C/MANIJA PALANCA
	4	4	6	4	9	0				C	BISAGRA DE ALUMINIO 3 1/2"x3 1/2" PESADA
x										C	PINTURA
	4	4	6	4	10	0				C	PINTURA LATEX 2 MANOS EN CIELO RASO
	4	4	6	4	10	0				C	PINTURA LATEX 2 MANOS EN MUROS INTERIORES Y EXTERIORES
										C	VARIOS
	4	4	6	4	11	0				C	JUNTA DE DILATACION CON ESPUMA PLASTICA + JEBE MICROPOROSO
	4	4	6	4	11	0				C	BARRA DE SEGURIDAD DE ALUMINIO EN BAÑOS P/DISCAPACITADOS
											OBRAS EXTERIORES 
											PORTICO DE INGRESO Y CERCO PERIMETRICO
										C	ESTRUCTURAS
x										C	MOVIMIENTO DE TIERRAS
	4	9	1	1	0	0				C	EXCAVACION MANUAL PARA CIMIENTOS EN TERRENO NORMAL
	4	9	1	1	0	0				C	RELLENO COMPACTADO CON MATERIAL PROPIO
	4	9	1	1	0	0				C	AFIRMADO COMPACTADO DE 8'' 
	4	9	1	1	0	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
x										C	CONCRETO SIMPLE
	4	9	1	2	0	0				C	CIMIENTOS CORRIDOS MEZCLA 1:10 CEMENTO:HORMIGON 30% PIEDRA
x										C	CONCRETO ARMADO
										C	ZAPATAS
	4	9	1	3	1	0				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	9	1	3	1	0				C	CONCRETO f'c=210 kg/cm2 EN ZAPATAS
										C	SOBRECIMIENTO
	4	9	1	3	2	0				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	9	1	3	2	0				C	ENCOFRADO Y DESENCOFRADO DE SOBRECIMIENTOS
	4	9	1	3	2	0				C	CONCRETO f'c=175 kg/cm2 PARA SOBRECIMIENTOS
										C	PLACAS
	4	9	1	3	3	0				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	9	1	3	3	0				C	ENCOFRADO Y DESENCOFRADO DE PLACAS DE CONCRETO
	4	9	1	3	3	0				C	CONCRETO f'c=210 kg/cm2 EN PLACAS
										C	VIGAS
	4	9	1	3	4	0				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	9	1	3	4	0				C	ENCOFRADO Y DESENCOFRADO DE VIGAS
	4	9	1	3	4	0				C	CONCRETO f'c=210 kg/cm2 EN VIGAS
										C	MURETE DE MEDIDOR
	4	9	1	3	4	0				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	9	1	3	5	0				C	ENCOFRADO Y DESENCOFRADO DE MURETE DE MEDIDOR
	4	9	1	3	5	0				C	CONCRETO f'c=210 kg/cm2 EN MURO DE MEDIDOR
										C	ARQUITECTURA Y ACABADOS
x										C	REVOQUES Y ENLUCIDOS
	4	9	6	1	0	0				C	TARRAJEO EN MUROS INTERIORES Y EXTERIORES CON C:A - 1:5
	4	9	6	1	0	0				C	REVESTIMIENTO CON PIEDRA GRANALLA COLOR ROJO
	4	9	6	1	0	0				C	REVESTIMIENTO CON GRANITO BLANCO
	4	9	6	1	0	0				C	REVESTIMIENTO EN MADERA CEDRO O SIMILAR
	4	9	6	1	0	0				C	REVESTIMIENTO EN PIEDRA LAJA GRIS 
x										C	CARPINTERIA METALICA
	4	9	6	2	0	0				C	PORTON DE INGRESO PRINCIPAL C/PERFILES METALICOS
	4	9	6	2	0	0				C	PUERTA DE INGRESO PRINCIPAL C/PERFILES METALICOS
	4	9	6	2	0	0				C	CERCO DE MALLA OLIMPICA 2"x2" C/PERFILES, SEGÚN DISEÑO
x										C	PINTURA
	4	9	6	3	0	0				C	PINTURA ESMALTE OLEO MATE 2 MANOS EN SOBRECIMIENTO DE CERCO METALICO
	4	9	6	3	0	0				C	PINTURA ESMALTE EPOXICO EN CARPINTERIA METALICA
	4	9	6	3	0	0				C	PINTURA LATEX 2 MANOS EN MUROS INTERIORES Y EXTERIORES
										C	VARIOS
	4	9	6	4	0	0				C	SUMINISTRO Y COLOCACION DE LETRAS DE ACERO INOXIDABLE ALTO RELIEVE "LOSA DEPORTIVA SACRAMENTO BAJO"
	4	9	6	4	0	0				C	SUMINISTRO E INSTALACION DE ESTRUCTURA DE ACERO PARA BASE DE LETRAS
	4	9	6	4	0	0				C	DISEÑO EN BAJO RELIEVE DE FIGURAS DEPORTIVAS, SEGÚN DISEÑO
	4	9	6	4	0	0				C	SUMINISTRO E INSTALACION DE PERFILES DE CAÑA EN FACHADA PRINCIPAL
										C	PISOS Y OBRAS EXTERIORES
										C	PISOS Y SARDINELES
x										C	MOVIMIENTO DE TIERRAS
	4	5	1	1	0	0				C	EXCAVACION MANUAL PARA SADINELES EN TERRENO NORMAL
	4	5	1	1	0	0				C	CORTE DE TERRENO A NIVEL DE SUB-RASANTE C/EQUIPO PESADO
	4	5	1	1	0	0				C	ESCARIFICADO Y COMPACTADO DE SUB-RASANTE C/ EQUIPO LIVIANO
	4	5	1	1	0	0				C	BASE GRANULAR COMPACTADA DE 4'' 
	4	5	1	1	0	0				C	BASE GRANULAR COMPACTADA DE 6'' 
	4	5	1	1	0	0				C	CAMA DE ARENA PARA ASENTADO DE ADOQUINES E=4cm
	4	9	1	1	0	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
x										C	CONCRETO SIMPLE
	4	5	1	2	0	0				C	ENCOFRADO Y DESENCOFRADO EN SARDINELES
	4	5	1	2	0	0				C	CONCRETO f'c=175 kg/cm2 PARA SARDINELES SUMERGIDOS
	4	5	1	2	0	0				C	CONCRETO f'c=175 kg/cm2 PARA PISOS, E=10cm.
	4	5	1	2	0	0				C	RAMPAS DE CONCRETO f'c=175 kg/cm2, e=10cm ACABADO FROTACHADO Y BRUÑADO
	4	5	1	2	0	0				C	VEREDA DE CONCRETO f'c=175 kg/cm2, e=10cm ACABADO SEMIPULIDO Y BRUÑADO
x										C	CONCRETO ARMADO
	4	5	1	3	0	0				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	5	1	3	0	0				C	ENCOFRADO Y DESENCOFRADO EN SARDINELES
	4	5	1	3	0	0				C	CONCRETO f'c=210kg/cm2 PARA SARDINELES REFORZADOS
										C	ARQUITECTURA Y ACABADOS
	4	5	1	4	0	0				C	PISO DE CONCRETO DISEÑO ESTAMPADO TIPO PIEDRA 
	4	5	1	4	0	0				C	PISO DE ADOQUIN COLOR AMARILLO 
	4	5	1	4	0	0				C	PISO DE ADOQUIN COLOR ROJO 
	4	5	1	4	0	0				C	PISO DE GRANITO PULIDO COLOR GRIS 
	4	5	1	4	0	0				C	LOSETA DE CAUCHO MULTICOLOR DE 50X50X30mm 
	4	5	1	4	0	0				C	SELLADO DE JUNTAS DE ADOQUIN CON ARENA 
	4	5	1	4	0	0				C	JUNTA DE DILATACION RELLENO CON MORTERO ASFALTICO E=1"
	4	5	1	4	0	0				C	SELLADO ACRILICO EPOXICO EN PISOS POROSOS
x										C	AREAS VERDES
	4	5	4	0	0	0				C	SUELO AGRICOLA PARA AREAS VERDES
	4	5	4	0	0	0				C	SEMBRADO DE GRASS NATURAL AMERICANO
	4	5	4	0	0	0				C	SUMINISTRO Y SEMBRIO DE ARBOLES Y ARBUSTOS
	4	5	4	0	0	0				C	JARDINERAS EXTERIORES CON BANCAS (INC. REVESTIMIENTO Y PINTURA)
										C	CONSTRUCCION DE OBELISCO
x										C	ESTRUCTURAS
	4	5	2	1	0	0				C	ACERO CORRUGADO FY=4200 kg/cm2 GRADO 60
	4	5	2	1	0	0				C	ENCOFRADO Y DESENCOFRADO EN OBELISCO
	4	5	2	1	0	0				C	CONCRETO f'c=210kg/cm2 EN OBELISCO
										C	ARQUITECTURA Y ACABADOS
	4	5	2	2	0	0				C	REVESTIMIENTO EN TERRAZO PULIDO COLOR BEIGE
x										C	CUBIERTA DE LONA DE PVC
	4	5	3	0	0	0				C	EXCAVACION MANUAL PARA DADOS DE CONCRETO
	4	5	3	0	0	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
	4	5	3	0	0	0				C	DADOS PARA COLUMNAS f'c=175 lg/cm2
	4	5	3	0	0	0				C	POSTE METALICO CIRCULAR DE 3" H:4.00m
	4	5	3	0	0	0				C	CUBIERTA DE LONA DE PVC
x										C	EQUIPAMIENTO Y MOBILIARIO URBANO
	4	5	5	0	0	0				C	BANCA CENTRAL CON JARDINERA CON ASIENTOS REVESTIDOS CON LISTONES DE MADERA Y GRANITO PULIDO COLOR BEIGE
	4	5	5	0	0	0				C	BANQUETA DE CONCRETO CON ACABADO EN GRANITO COLOR BEIGE Y TABLONES DE MADERA
	4	5	5	0	0	0				C	BANCAS Y TABLERO DE CONCRETO PARA JUEGOS DE AJEDREZ REVESTIDAS CON GRANALLA DE COLOR SEGÚN DISEÑO
	4	5	5	0	0	0				C	INSTALACION DE SOL Y SOMBRA DE ALUMINIO NEGRO SEGÚN DISEÑO,  PARA BANCAS CENTRALES 
	4	5	5	0	0	0				C	INSTALACION DE SOL Y SOMBRA DE ALUMINIO NEGRO SEGÚN DISEÑO, PARA BANQUETAS DE CONCRETO Y JUEGOS DE AJEDREZ
	4	5	5	0	0	0				C	SUMINISTRO E INSTALACION DE TACHOS DE BASURA (SEGÚN DISEÑO)
x										C	JUEGOS INFANTILES
	4	5	6	0	0	0				C	SUMINISTRO E INSTALACION DE JUEGOS INFANTILES SEGÚN PLANOS
										C	VARIOS
	4	5	7	0	0	0				C	SUMINISTRO E INSTALACION DE MAQUINAS PARA ZONAS DE GYM
	4	5	7	0	0	0				C	PLACA RECORDATORIA
										C	INSTALACIONES SANITARIAS
										C	SISTEMA DE AGUA POTABLE
x										C	APARATOS SANITARIOS Y ACCESORIOS
	4	6	1	1	0	0				C	INODORO TANQUE BAJO COLOR BLANCO
	4	6	1	1	0	0				C	LAVATORIO OVALIN INC. ACCESORIOS
	4	6	1	1	0	0				C	LAVATORIO BLANCO CON PEDESTAL INC. GRIFERIA
	4	6	1	1	0	0				C	URINARIO CON FLUXOMETRO INC. ACCESORIOS
	4	6	1	1	0	0				C	PORTA PAPEL DE LOSA BLANCA DE SOBREPONER
	4	6	1	1	0	0				C	SALIDA PARA DUCHA
	4	6	1	1	0	0				C	CISTERNA DE AGUA 1200 lt. INC. ACCESORIOS
	4	6	1	1	0	0				C	TANQUE ELEVADO DE 1100 lt. INC. ACCESORIOS
										C	SISTEMA DE AGUA FRIA
	4	6	1	2	0	0				C	EXCAVACION DE ZANJAS PARA RED DE AGUA FRIA
	4	6	1	2	0	0				C	REFINE Y NIVELACION DE ZANJAS PARA TUBERIA
	4	6	1	2	0	0				C	CAMA DE ARENA GRUESA PARA RED DE TUBERIA DE AGUA FRIA (e=0.10 mts).
	4	6	1	2	0	0				C	RELLENO MANUAL CON MATERIAL PROPIO (REDES DE AGUA FRIA)
	4	6	1	2	0	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
	4	6	1	2	0	0				C	TUBERIA DE PVC SAP 1/2"
	4	6	1	2	0	0				C	TUBERIA DE PVC SAP 3/4"
	4	6	1	2	0	0				C	TUBERIA DE PVC SAP 1"
	4	6	1	2	0	0				C	CODO PVC SAP 1/2"
	4	6	1	2	0	0				C	CODO PVC SAP 3/4"
	4	6	1	2	0	0				C	CODO PVC SAP 1"
	4	6	1	2	0	0				C	TEE PVC SAP 1/2"
	4	6	1	2	0	0				C	TEE PVC SAP 3/4"
	4	6	1	2	0	0				C	REDUCCION PVC SAP C10 3/4" A 1/2"
	4	6	1	2	0	0				C	VALVULA ESFERICA DE BRONCE DE 3/4" DE UNION ROSCADA
	4	6	1	2	0	0				C	VALVULA COMPUERTA DE BRONCE DE UNION ROSCADA DE 3/4"
	4	6	1	2	0	0				C	VALVULA COMPUERTA DE BRONCE DE UNION ROSCADA DE 1"
	4	6	1	2	0	0				C	VALVULA CHECK DE BRONCE 3/4"
	4	6	1	2	0	0				C	VALVULA CHECK DE BRONCE 1"
	4	6	1	2	0	0				C	CAJA PARA VALVULA EN PARED, NICHO DE MAYOLICA Y TAPA METALICA
	4	6	1	2	0	0				C	LLAVE DE RIEGO C/GRIFO DE 1/2" EN CAJUELA DE CONCRETO 
	4	6	1	2	0	0				C	CAJA DE REGISTRO PARA VALVULA DE AGUA
x										C	SISTEMA DE DESAGUE
										C	RED EXTERIOR DE DESAGUIE
	4	6	3	1	0	0				C	EXCAVACIONES A PULSO TERRENO NORMAL PARA TUBERIA 2" - 4" HASTA 1.00m DE PROFUNDIDAD
	4	6	3	1	0	0				C	REFINE Y NIVELACION ZANJA TERRENO CONGLOMERADO PARA TUBERIA 2" - 4"
	4	6	3	1	0	0				C	RELLENO COMPACTADO DE ZANJAS HASTA 1.00m
	4	6	3	1	0	0				C	ELIMINACION DE MATERIAL EXCEDENTE D. PROM=15.8 KM
	4	6	3	1	0	0				C	TUBERIA DE PVC SAL 4"
	4	6	3	1	0	0				C	CAJA DE REGISTRO DE DESAGUE 12" x 24" 
										C	RED INTERIOR DE DESAGUE
	4	6	3	2	0	0				C	SALIDA DE DESAGUE DE PVC 2"
	4	6	3	2	0	0				C	SALIDA DE DESAGUE DE PVC 4"
	4	6	3	2	0	0				C	SALIDA DE VENTILACION DE PVC 2"
	4	6	3	2	0	0				C	TUBERIA DE PVC SAL 4"
	4	6	3	2	0	0				C	TUBERIA DE PVC SAL 2"
	4	6	3	2	0	0				C	REGISTRO DE BRONCE DE 4"
	4	6	3	2	0	0				C	YEE PVC SAL 4"
	4	6	3	2	0	0				C	YEE PVC SAL 4" x 2"
	4	6	3	2	0	0				C	CODO PVC SAL 90° DE 4" A 2"
	4	6	3	2	0	0				C	CODO PVC SAL 4" x 45°
	4	6	3	2	0	0				C	CODO PVC SAL SP DE 2" x 90°
	4	6	3	2	0	0				C	REDUCCION DE PVC SAL DE 4" A 2"
	4	6	3	2	0	0				C	VALVULA DE BOLA PVC DE 2"
	4	6	3	2	0	0				C	SOMBRERO DE VENTILACION PVC DE 2"
										C	SISTEMA DE DRENAJE PLUVIAL
	4	6	2	0	0	0				C	TUBERÍA DE DESAGÜE PVC-SAP Ø=4''
	4	6	2	0	0	0				C	CANALETA DE PLANCHA GALVANIZADA  E=0.40MM, PARA COBERTURA METALICA
	4	6	2	0	0	0				C	CUNETA 1/2 CAÑA EN VIGA ACANALADA ACAB. 1:5 C:A
	4	6	2	0	0	0				C	EMBUDO METALICO 4" PARA CAIDA
	4	6	2	0	0	0				C	ABRAZADERA DE FIJACION F°G° 4"
										C	PRUEBAS HIDRAULICAS
	4	6	4	0	0	0				C	PRUEBAS HIDRAULICAS DE RED DE AGUA POTABLE
	4	6	4	0	0	0				C	PRUEBAS HIDRAULICAS DE RED DE DESAGUE
x										C	INSTALACIONES ELECTRICAS
										C	SALIDAS PARA ALUMBRADO, TOMACORRIENTE Y FUERZA
										C	SALIDAS PARA ALUMBRADO
	4	7	1	1	0	0				C	SALIDA PARA ALUMBRADO EN TECHO
	4	7	1	1	0	0				C	SALIDA PARA ALUMBRADO BRAQUET 
	4	7	1	1	0	0				C	SALIDA PARA ALUMBRADO EN PISO
										C	SALIDAS PARA INTERRUPTORES
	4	7	1	2	0	0				C	SALIDA PARA INTERRUPTOR UNIPOLAR SIMPLE
										C	SALIDAS PARA TOMACORRIENTES
	4	7	1	3	0	0				C	TOMACORRIENTE BIPOLAR DOBLE CON LINEA A TIERRA
										C	SALIDAS DE FUERZA
	4	7	1	4	0	0				C	SALIDA DE FUERZA PARA ELECTROBOMBA DE 3/4 HP
										C	SISTEMA DE CONDUCTOS
x										C	MOVIMIENTO DE TIERRAS
	4	7	2	1	0	0				C	EXCAVACION DE ZANJAS PARA REDES ELECTRICAS
	4	7	2	1	0	0				C	EXCAVACION DE ZANJAS PARA POSTES
	4	7	2	1	0	0				C	EXCAVACION MANUAL PARA BUZON ENERGIA H=0.65M 
	4	7	2	1	0	0				C	RELLENO MANUAL C/MATERIAL PROPIO EN REDES ELECTRICAS 
	4	7	2	1	0	0				C	CAMA DE ARENA H=0.10M 
	4	7	2	1	0	0				C	CINTA DE SEÑALIZACION DE CABLE SUBTERRANEO 
	4	7	2	1	0	0				C	ELIMINACION DE MATERIAL PROVENIENTE DE EXCAVACONES (D. PROM. =10km)
										C	ELECTRODUCTOS
	4	7	2	2	0	0				C	TUBERIA PVC-P (ELECTRICAS) D=25 mm 
	4	7	2	2	0	0				C	TUBERIA PVC-P (ELECTRICAS) D=35 mm 
										C	CONDUCTORES
	4	7	2	3	0	0				C	CABLE NH-80 (2-1x2.5mm2+1x2.5mm2(T)) 
	4	7	2	3	0	0				C	CABLE NH-80 (2-1x4mm2+1x4mm2(T)) 
	4	7	2	3	0	0				C	CABLE N2XOH (2-1x4mm2+1x4mm2(T)) 
	4	7	2	3	0	0				C	CABLE N2XOH (2-1x6mm2+1x6mm2(T)) 
	4	7	2	3	0	0				C	CABLE N2XOH (2-1x10mm2+1x10mm2(T)) 
										C	BUZONES
	4	7	2	4	0	0				C	BUZON ELECTRICO, CONCRETO f’´c=210kg/cm2, H=0.65M 
										C	TABLEROS
	4	7	2	5	0	0				C	TABLERO GENERAL AUTOSOPORTADO - TG
	4	7	2	5	0	0				C	TABLERO DE DISTRIBUCION STD-1
	4	7	2	5	0	0				C	TABLERO DE DISTRIBUCION STD-2
										C	VARIOS
	4	7	2	6	0	0				C	SUMINISTRO E INSTALACION DE ELECTROBOMBA MONOFASICO DE 3/4 HP, INCLUIDO ACCESORIOS
	4	7	2	6	0	0				C	SUMINISTRO DE ENERGIA ELECTRICA (Inc. Reubicacion de medidor y tramite de suministro de demanda)
										C	ARTEFACTOS DE ALUMBRADO
	4	7	3	0	0	0				C	LUMINARIA PANEL LED CIRCULAR 24W ADOSABLE LUZ BLANCA
	4	7	3	0	0	0				C	LUMINARIA EN PISO 20W
	4	7	3	0	0	0				C	EQUIPO PORTALAMPARAS WALL PACK LED 45W O SIMILAR. 
	4	7	3	0	0	0				C	LUMINARIA TIPO SPOTH LIGTH EN PARED
	4	7	3	0	0	0				C	LUMINARIA LED TIPO FAROLA 50W EN TUBERIA DE 6m de F°G°(D=60mm), 3000K
	4	7	3	0	0	0				C	LUMINARIA REFLECTOR LED DE 200W - LUZ BLANCA  
										C	SISTEMA DE PUESTA A TIERRA
	4	7	4	0	0	0				C	POZO DE TIERRA (R<10 Ohms)
	4	7	4	0	0	0				C	CABLE DESNUDO Cu 35mm2 
										C	PRUEBAS ELECTRICAS
	4	7	5	0	0	0				C	PRUEBA DE AISLAMIENTO
	4	7	5	0	0	0				C	PRUEBA DE ATERRAMIENTO
										C	OTROS
	4	8	0	0	0	0				C	MITIGACION DE IMPACTO AMBIENTAL
	4	8	0	0	0	0				C	ELABORACION DE PLAN DE MONITOREO ARQUEOLOGICO 
        """

        lines = [line for line in raw_text.split('\n') if line.strip()]

        parsed_packages = []
        current_package = None

        # Regex para identificar nombres de actividad (después de "C " o al inicio)
        # Busca patrones tipo "C " seguido de texto, ignorando números iniciales
        activity_regex = re.compile(r'(?:C\s+)(.+)$', re.IGNORECASE)

        for line in lines:
            line = line.strip()

            # Detectar si es un Paquete Principal (Empieza con 'x' o 'X')
            if line.lower().startswith('x'):
                # Si ya teníamos uno, lo guardamos
                if current_package:
                    parsed_packages.append(current_package)

                # Limpiar el nombre (quitar 'x', 'C', tabs)
                # Ejemplo: "x C MOVIMIENTO..." -> "MOVIMIENTO..."
                raw_name = re.sub(r'^[xC\s\t]+', '', line).strip()

                current_package = {
                    "name": raw_name,
                    "activities": []  # Lista para guardar las partidas
                }

            # Si no es paquete, asumimos que es actividad DEL paquete actual
            elif current_package:
                # Intentar extraer nombre limpio de la actividad
                match = activity_regex.search(line)
                if match:
                    clean_activity = match.group(1).strip()
                    current_package["activities"].append(clean_activity)
                else:
                    # Fallback: si no hay "C ", limpiar números iniciales y tabulaciones
                    # Ejemplo: "4 1 0 ... CEMENTO" -> "CEMENTO"
                    # Regex busca la primera letra mayuscula despues de numeros/espacios
                    clean_match = re.search(r'[A-Z][A-Z\s\(\)\.\,\d\+\-]+$', line)
                    if clean_match:
                        current_package["activities"].append(clean_match.group(0).strip())

        # Agregar el último
        if current_package:
            parsed_packages.append(current_package)

        logger.info(f"Se detectaron {len(parsed_packages)} paquetes de trabajo principales.")
        return parsed_packages

    def get_description_for_package(self, package_name):
        """Busca una descripción inteligente basada en palabras clave."""
        # Búsqueda exacta primero
        if package_name in self.knowledge_base:
            return self.knowledge_base[package_name]

        # Búsqueda parcial (ej. "COBERTURA METALICA" coincidirá con "COBERTURA")
        for key, desc in self.knowledge_base.items():
            if key in package_name:
                return desc

        return "Ejecución de la partida de acuerdo a las especificaciones técnicas del proyecto, planos de detalle y normas vigentes. Incluye mano de obra, materiales y equipos necesarios."

    def get_project_meta(self):
        """Datos fijos del proyecto."""
        return {
            "version": "1.0",
            "proyecto": "Mejoramiento de los Servicios Públicos de Integración Económica y Social en la Losa Deportiva del A.H. Sacramento Bajo – Palpa",
            "preparado_por": "Ing. Residente",
            "revisado_por": "Ing. Supervisor",
            "aprobado_por": "Gerencia de Infraestructura",
            "fecha": datetime.now().strftime("%d/%m/%Y")
        }

    # --- UTILITARIOS DE ESTILO WORD ---
    def _set_cell_bg(self, cell, color_hex):
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _format_cell(self, cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=8, bg_color=None):
        cell.text = str(text)
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.runs[0]
        run.font.bold = bold
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
        if bg_color:
            self._set_cell_bg(cell, bg_color)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(
            r'<w:tcMar {}><w:top w:w="20" w:type="dxa"/><w:bottom w:w="20" w:type="dxa"/></w:tcMar>'.format(
                nsdecls('w')))
        tcPr.append(tcMar)

    def _create_page_for_item(self, document, wbs_data, item_index, meta):
        """
        Genera UNA página completa de diccionario para un ítem específico.
        wbs_data: Diccionario con 'name' y 'activities'
        """
        wbs_name = wbs_data['name']
        activities = wbs_data['activities']

        f_date = meta['fecha'].split('/')
        description = self.get_description_for_package(wbs_name)
        wbs_id = f"1.{item_index + 1}"

        # ---------------------------------------------------------
        # 1. TABLA DE CABECERA (Header)
        # ---------------------------------------------------------
        table = document.add_table(rows=6, cols=6)
        table.style = 'Table Grid'

        # Fila 0: Título
        row0 = table.rows[0]
        c = row0.cells[0].merge(row0.cells[5])
        self._format_cell(c, "DICCIONARIO DE LA EDT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)

        # Fila 1: Versión
        row1 = table.rows[1]
        c = row1.cells[0].merge(row1.cells[5])
        self._format_cell(c, f"Versión {meta['version']}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Fila 2: Proyecto
        row2 = table.rows[2]
        self._format_cell(row2.cells[0], "PROYECTO")
        c = row2.cells[1].merge(row2.cells[5])
        self._format_cell(c, meta['proyecto'])

        # Firmas
        def add_signature_row(row_idx, label, name):
            r = table.rows[row_idx]
            self._format_cell(r.cells[0], label)
            self._format_cell(r.cells[1], name)
            self._format_cell(r.cells[2], "FECHA")
            self._format_cell(r.cells[3], f_date[0], align=WD_ALIGN_PARAGRAPH.CENTER)
            self._format_cell(r.cells[4], f_date[1], align=WD_ALIGN_PARAGRAPH.CENTER)
            self._format_cell(r.cells[5], f_date[2], align=WD_ALIGN_PARAGRAPH.CENTER)

        add_signature_row(3, "PREPARADO POR:", meta['preparado_por'])
        add_signature_row(4, "REVISADO POR:", meta['revisado_por'])
        add_signature_row(5, "APROBADO POR:", meta['aprobado_por'])

        for row in table.rows:
            row.cells[0].width = Cm(3.5)

        document.add_paragraph()

        # ---------------------------------------------------------
        # 2. PAQUETE DE TRABAJO
        # ---------------------------------------------------------
        t_wbs = document.add_table(rows=1, cols=4)
        t_wbs.style = 'Table Grid'
        self._format_cell(t_wbs.rows[0].cells[0], "Paquete de trabajo :", bold=True)
        self._format_cell(t_wbs.rows[0].cells[1], wbs_name, bold=True, bg_color="FFFF99")
        self._format_cell(t_wbs.rows[0].cells[2], "ID E.D.T. :")
        self._format_cell(t_wbs.rows[0].cells[3], wbs_id)

        t_wbs.rows[0].cells[0].width = Cm(3.5)
        t_wbs.rows[0].cells[1].width = Cm(8.0)
        t_wbs.rows[0].cells[2].width = Cm(2.0)

        # ---------------------------------------------------------
        # 3. DESCRIPCIÓN
        # ---------------------------------------------------------
        t_desc = document.add_table(rows=1, cols=1)
        t_desc.style = 'Table Grid'
        desc_cell = t_desc.rows[0].cells[0]
        p = desc_cell.paragraphs[0]
        run_title = p.add_run("Descripción del trabajo.- ")
        run_title.font.bold = True
        run_title.font.size = Pt(8)
        run_title.font.name = 'Arial'
        run_text = p.add_run(description)
        run_text.font.size = Pt(8)
        run_text.font.name = 'Arial'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ---------------------------------------------------------
        # 4. HITOS (Genéricos)
        # ---------------------------------------------------------
        t_milestones = document.add_table(rows=1, cols=2)
        t_milestones.style = 'Table Grid'

        c_left = t_milestones.rows[0].cells[0]
        p_left = c_left.paragraphs[0]
        p_left.add_run("Hitos\n").font.bold = True
        p_left.add_run("\t1. Inicio de Actividad\n")
        p_left.add_run("\t2. Fin de Actividad\n")

        c_right = t_milestones.rows[0].cells[1]
        p_right = c_right.paragraphs[0]
        p_right.add_run("Fechas:\n").font.bold = True
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.add_run("Según Cronograma\n")
        p_right.add_run("Según Cronograma\n")

        # ---------------------------------------------------------
        # 5. TABLA DE COSTOS / ACTIVIDADES
        # ---------------------------------------------------------
        # Calcular filas necesarias (Header + Actividades + Relleno si faltan)
        num_activities = len(activities)
        min_rows = 5  # Mínimo de filas visuales

        # Crear tabla
        t_costs = document.add_table(rows=2, cols=10)
        t_costs.style = 'Table Grid'

        # Headers
        r0 = t_costs.rows[0]
        r1 = t_costs.rows[1]

        c_id = r0.cells[0].merge(r1.cells[0])
        self._format_cell(c_id, "ID", align=WD_ALIGN_PARAGRAPH.CENTER)
        c_act = r0.cells[1].merge(r1.cells[1])
        self._format_cell(c_act, "Actividad", align=WD_ALIGN_PARAGRAPH.CENTER)
        c_res = r0.cells[2].merge(r1.cells[2])
        self._format_cell(c_res, "Recursos", align=WD_ALIGN_PARAGRAPH.CENTER)
        c_hh = r0.cells[3].merge(r0.cells[5])
        self._format_cell(c_hh, "Horas Hombre", align=WD_ALIGN_PARAGRAPH.CENTER)
        c_mat = r0.cells[6].merge(r0.cells[8])
        self._format_cell(c_mat, "Materiales", align=WD_ALIGN_PARAGRAPH.CENTER)
        c_tot = r0.cells[9].merge(r1.cells[9])
        self._format_cell(c_tot, "Costo Total", align=WD_ALIGN_PARAGRAPH.CENTER)

        subheaders = ["Horas", "Ratio", "Total", "Unidades", "Costo", "Total"]
        for i, text in enumerate(subheaders):
            self._format_cell(r1.cells[3 + i], text, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Llenar con actividades reales del parsing
        for idx, activity in enumerate(activities):
            row = t_costs.add_row()
            cells = row.cells
            # ID actividad relativo
            self._format_cell(cells[0], f"{wbs_id}.{idx + 1}", align=WD_ALIGN_PARAGRAPH.CENTER)
            # Nombre de la actividad
            self._format_cell(cells[1], activity, align=WD_ALIGN_PARAGRAPH.LEFT)
            # Celdas vacías para el resto
            for k in range(2, 10):
                self._format_cell(cells[k], " ")

        # Rellenar con filas vacías si hay pocas actividades (para que no quede fea la tabla)
        rows_to_add = max(0, min_rows - num_activities)
        for _ in range(rows_to_add):
            row = t_costs.add_row()
            for cell in row.cells:
                self._format_cell(cell, " ")

        # ---------------------------------------------------------
        # 6. FOOTER
        # ---------------------------------------------------------
        t_foot = document.add_table(rows=4, cols=1)
        t_foot.style = 'Table Grid'

        def add_footer_row(idx, label, text, underline=False):
            cell = t_foot.rows[idx].cells[0]
            p = cell.paragraphs[0]
            run_lbl = p.add_run(label + " ")
            run_lbl.font.bold = True
            run_lbl.font.size = Pt(8)
            run_val = p.add_run(text)
            run_val.font.size = Pt(8)
            if underline:
                run_val.font.underline = True

        add_footer_row(0, "Requisitos de Calidad:",
                       "Control de trazos y niveles, pruebas de calidad (slump, probetas) según RNE.")
        add_footer_row(1, "Criterios de aceptación:",
                       "Aprobación del Supervisor de Obra según especificaciones técnicas.", underline=True)
        add_footer_row(2, "Información Técnica:", "Planos de Arquitectura, Estructuras e Instalaciones.")
        add_footer_row(3, "Información Contractual:", "Precios Unitarios Contractuales")

    def create_report(self):
        """Orquesta la creación y guardado del documento masivo."""
        try:
            logger.info("Iniciando generación de reporte EDT Masivo...")

            # 1. Obtener datos parseados (Estructura de Objetos)
            meta = self.get_project_meta()
            wbs_packages = self.parse_complex_input()

            if not wbs_packages:
                logger.warning("No se detectaron paquetes. Verifica el formato del input.")
                return

            # 2. Configurar documento
            document = Document()
            section = document.sections[0]
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)

            # 3. Generar páginas
            total_items = len(wbs_packages)
            for i, pkg_data in enumerate(wbs_packages):
                logger.info(f"Generando página {i + 1}/{total_items}: {pkg_data['name']}")
                self._create_page_for_item(document, pkg_data, i, meta)

                # Agregar salto de página excepto en el último ítem
                if i < total_items - 1:
                    document.add_page_break()

            # 4. Guardar
            safe_project_name = "Diccionario_EDT_Losa_Sacramento_Bajo_Completo"
            filename = f"{safe_project_name}.docx"
            file_path = self.output_dir / filename

            try:
                document.save(file_path)
            except PermissionError:
                logger.error(f"El archivo {filename} está abierto. Ciérralo.")
                return

            logger.info(f"Archivo generado exitosamente: {file_path}")
            os.startfile(file_path)

        except Exception as e:
            logger.error(f"Error crítico en el proceso: {e}")
            raise


if __name__ == "__main__":
    generator = WBSReportGenerator()
    generator.create_report()