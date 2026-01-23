import logging
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Configuración de Logging ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("generacion_word.log"),
        logging.StreamHandler()
    ]
)


def set_cell_border(cell, **kwargs):
    """
    Función auxiliar para bordes de celda en python-docx.
    Esto es necesario porque la librería estándar no expone bordes de celdas fácilmente.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        if border_name in kwargs:
            element = OxmlElement(f'w:{border_name}')
            element.set(qn('w:val'), kwargs[border_name])
            element.set(qn('w:sz'), '4')  # Tamaño de la línea
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), 'auto')
            tcPr.append(element)


def crear_documento_siga():
    logging.info("Iniciando generación del documento Word replicado...")

    try:
        document = Document()

        # Configuración de márgenes (A4 estándar aproximado)
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

        # --- ESTILOS ---
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(8)  # SIGA usa letra pequeña generalmente

        # --- ENCABEZADO ---
        header_p = document.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run("Sistema Integrado de Gestión Administrativa\nMódulo de Logística\nVersión 25.01.02")
        run.bold = True
        run.font.size = Pt(9)

        # --- TÍTULO ---
        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = title_p.add_run("\nSOLICITUD DE MODIFICACIÓN DEL CUADRO MULTIANUAL DE NECESIDADES N° 0000001920\n")
        run_title.bold = True
        run_title.font.size = Pt(11)

        # --- DATOS GENERALES (Tabla invisible para layout) ---
        table_info = document.add_table(rows=1, cols=2)
        table_info.autofit = True

        # Columna Izquierda (Datos Entidad)
        cell_left = table_info.cell(0, 0)
        p_left = cell_left.paragraphs[0]
        p_left.add_run("UNIDAD EJECUTORA: 001 MUNICIPALIDAD PROVINCIAL DE YUNGUYO\n").bold = True
        p_left.add_run("NRO. IDENTIFICACIÓN: 301694\n")
        p_left.add_run("Centro de Costo: 030605 CREACION DE LOS SERVICIOS CULTURALES TEATRO MUNICIPAL\n")
        p_left.add_run("Fecha de Solicitud: 21/11/2025")

        # Columna Derecha (Fecha Impresión)
        cell_right = table_info.cell(0, 1)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.add_run("Fecha: 21/11/2025\n")
        p_right.add_run("Hora: 16:26:34\n")
        p_right.add_run("Página: 1 de 1")

        document.add_paragraph()  # Espacio

        # --- TABLA DE ITEMS ---
        # Definimos la estructura: 7 columnas
        # Headers complejos: Fila 1 y Fila 2
        table = document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        # Encabezados Fila 1
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Código Item N."
        hdr_cells[1].text = "Descripción del Item"
        hdr_cells[2].text = "Unidad de Medida"
        hdr_cells[3].text = "EXCLUSIÓN"
        hdr_cells[4].text = ""  # Placeholder para merge
        hdr_cells[5].text = "INCLUSIÓN"
        hdr_cells[6].text = ""  # Placeholder para merge

        # Merge de celdas de encabezado principal
        hdr_cells[3].merge(hdr_cells[4])
        hdr_cells[5].merge(hdr_cells[6])

        # Fila 2 de encabezados (Subtítulos)
        row2 = table.add_row().cells
        row2[0].text = ""
        row2[1].text = ""
        row2[2].text = ""
        row2[3].text = "Cantidad Total"
        row2[4].text = "Valor Total S/"
        row2[5].text = "Cantidad Total"
        row2[6].text = "Valor Total S/"

        # Datos extraídos del PDF
        data = [
            ["283400480026", "TABLERO ELECTRICO DE METAL PARA EMPOTRAR DE 24 POLOS INTERRUPTOR RIEL DIN", "Unidad",
             "0.00", "0.00", "3.00", "0.00"],
            ["283400480036", "TABLERO ELECTRICO DE METAL PARA ADOSAR DE 32 POLOS INTERRUPTOR RIEL DIN", "Unidad",
             "0.00", "0.00", "1.00", "0.00"],
            ["283400480096", "TABLERO ELECTRICO DE METAL PARA EMPOTRAR DE 30 POLOS INTERRUPTOR RIEL DIN", "Unidad",
             "0.00", "0.00", "2.00", "0.00"],
            ["283400480125",
             "TABLERO ELECTRICO DE METAL PARA EMPOTRAR DE 18 POLOS 10 KW 380/220 V 3 X 35 A INTERRUPTOR RIEL DIN",
             "Unidad", "0.00", "0.00", "5.00", "0.00"],
            ["283400480128",
             "TABLERO ELECTRICO DE METAL PARA EMPOTRAR DE 12 POLOS 7.5 KW 380/220 V 3 X 25 A INTERRUPTOR RIEL DIN",
             "Unidad", "0.00", "0.00", "2.00", "0.00"]
            # Nota: En el ítem final, asumí Cantidad=2.00 basado en lógica de compra, aunque el PDF era ambiguo visualmente.
        ]

        for item in data:
            row_cells = table.add_row().cells
            for i, text in enumerate(item):
                row_cells[i].text = text
                if i >= 3:  # Alinear números a la derecha
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph()  # Espacio

        # --- SUSTENTO Y NOTAS ---
        p_sustento = document.add_paragraph()
        run_s = p_sustento.add_run(
            "Sustento para la aprobación de modificaciones del CMN, al día hábil siguiente de su presentación (numeral 32.7 del articulo 32 de la Directiva): ")
        run_s.bold = True
        p_sustento.add_run(
            "ADQUICION DE BIENES TABLEROS PARA INSTALACIONES ELECTRICAS, PARA EL PROYECTO: CREACION DE LOS SERVICIOS CULTURALES PARA LA PARTICIPACION DE LA POBLACION EN LAS INDUSTRIAS CULTURALES Y LAS ARTE EN TEATRO MUNICIPAL DEL DISTRITO DE YUNGUYO")

        document.add_paragraph(
            "De ser el caso, indicar el/los año(s) que corresponda(n) realizar la inclusión o exclusión de la programación:")

        # Notas al pie
        notes = [
            "1/ La información registrada en el presente Anexo corresponde a campos mínimos y obligatorios que pueden ser ampliados por la Entidad del Sector Público u organización de la entidad.",
            "2/ La información registrada en los campos de 'exclusión' e 'inclusión' considera la cantidad y/o valor acumulado de todos los años de la programación.",
            "3/ El campo de 'cantidad total' se completa solo en el caso de bienes.",
            "4/ La presente información tiene carácter de Declaración Jurada; por lo que, en señal de conformidad y en representación del Área usuaria, se suscribe:"
        ]

        for note in notes:
            p_note = document.add_paragraph(note)
            p_note.style.font.size = Pt(7)

        document.add_paragraph("\n\n\n")  # Espacio para firmas

        # --- FIRMAS ---
        table_signatures = document.add_table(rows=1, cols=3)
        table_signatures.autofit = True

        # Firma Central (Usuario)
        sig_cell = table_signatures.cell(0, 1)
        p_sig = sig_cell.paragraphs[0]
        p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sig.add_run("__________________________\n")
        p_sig.add_run("Thon Borda Cealla\n").bold = True
        p_sig.add_run("CIP. N 57185\n")
        p_sig.add_run("RESIDENTE DE OBRA\n")
        p_sig.add_run("Firma: Responsable del Área Usuaria")

        # Guardar archivo
        output_filename = "Solicitud_CMN_Replicada.docx"
        document.save(output_filename)
        logging.info(f"Documento generado exitosamente: {output_filename}")
        print(f"¡Éxito! El archivo '{output_filename}' ha sido generado.")

    except Exception as e:
        logging.error(f"Error crítico generando el documento: {e}")
        print(f"Error: {e}")


if __name__ == "__main__":
    crear_documento_siga()